import os
import io
import json
import re
import traceback
import pandas as pd
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime, timedelta
from dotenv import load_dotenv
from pinecone import Pinecone
from flask import Flask, request, jsonify, send_file
from google.cloud import storage
import vertexai
from vertexai.language_models import TextEmbeddingModel
import requests
from docx import Document
from fpdf import FPDF

# --- CONFIGURATION ---
load_dotenv()
PROJECT_ID = os.environ.get("GOOGLE_CLOUD_PROJECT")
LOCATION = "asia-south1"
BUCKET_NAME = "rag-source-documents"
METADATA_FILE_PATH = "vector-search-inputs/metadata_lookup.json"
PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY")
PINECONE_INDEX_NAME = "tata-strive-rag"
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"
EMBEDDING_MODEL = "text-embedding-005"
RAG_MODEL = "google/gemini-flash-1.5-8b"
TRANSLATION_MODEL = "google/gemini-flash-1.5-8b"

# Email Configuration
EMAIL_SENDER = os.environ.get("EMAIL_SENDER_ADDRESS")
EMAIL_PASSWORD = os.environ.get("EMAIL_SENDER_PASSWORD")
EMAIL_SENDER_NAME = os.environ.get("EMAIL_SENDER_NAME", "Tata Strive Learning Team")
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587

# --- IN-MEMORY DATA STORAGE ---
course_data = {}  # Will store course duration info
holiday_data = {}  # Will store holidays by state
assessment_guidelines = ""  # Will store assessment guidelines

# --- CLIENT INITIALIZATION ---
vertexai.init(project=PROJECT_ID, location=LOCATION)
storage_client = storage.Client()
embedding_model = TextEmbeddingModel.from_pretrained(EMBEDDING_MODEL)
app = Flask(__name__)

print("Initializing Pinecone...")
pc = Pinecone(api_key=PINECONE_API_KEY)
pinecone_index = pc.Index(PINECONE_INDEX_NAME)
print("✅ Pinecone index connection established.")

print(f"Loading metadata from gs://{BUCKET_NAME}/{METADATA_FILE_PATH}...")
try:
    bucket = storage_client.bucket(BUCKET_NAME)
    blob = bucket.blob(METADATA_FILE_PATH)
    metadata_lookup = json.loads(blob.download_as_string())
    print("✅ Metadata loaded successfully.")
except Exception as e:
    print(f"❌ FATAL: Could not load metadata file. Error: {e}")
    metadata_lookup = {}

# --- LOCATION DETECTION ---
def detect_location_from_ip(ip_address: str) -> dict:
    """Detect location from IP address"""
    try:
        response = requests.get(f"http://ip-api.com/json/{ip_address}")
        data = response.json()
        return {
            "city": data.get("city", ""),
            "state": data.get("regionName", ""),
            "country": data.get("country", "India"),
            "detected": True
        }
    except:
        return {"city": "", "state": "", "country": "India", "detected": False}

def get_suggested_language(state: str) -> str:
    """Suggest language based on state"""
    language_map = {
        "West Bengal": "Bengali",
        "Maharashtra": "Marathi",
        "Gujarat": "Gujarati",
        "Tamil Nadu": "Tamil",
        "Karnataka": "Kannada",
        "Kerala": "Malayalam",
        "Andhra Pradesh": "Telugu",
        "Telangana": "Telugu",
        "Odisha": "Odia",
        "Punjab": "Punjabi",
        "Haryana": "Hindi",
        "Rajasthan": "Hindi",
        "Uttar Pradesh": "Hindi",
        "Madhya Pradesh": "Hindi",
        "Bihar": "Hindi"
    }
    return language_map.get(state, "English")

# --- DATA LOADING FUNCTIONS ---
def load_course_data(file):
    """Load course duration data from CSV"""
    global course_data
    try:
        df = pd.read_csv(file)
        for _, row in df.iterrows():
            course_data[row['name']] = {
                'id': row['id'],
                'duration_hours': row.get('cumulative_course_duration', 0),
                'theory_hours': row.get('domain_theory_hours', 0),
                'eligibility': row.get('eligibility_criteria', 'Not specified')
            }
        print(f"✅ Loaded {len(course_data)} courses")
        return {"success": True, "courses_loaded": len(course_data)}
    except Exception as e:
        print(f"❌ Error loading course data: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def load_holiday_data(file):
    """Load holiday data from CSV"""
    global holiday_data
    try:
        df = pd.read_csv(file)
        for state in df['Location'].unique():
            state_holidays = df[df['Location'] == state]
            holiday_data[state] = []
            for _, row in state_holidays.iterrows():
                try:
                    holiday_date = datetime.strptime(row['HolidayDate'], '%d-%m-%Y')
                    holiday_data[state].append({
                        'name': row['Holidays'],
                        'date': holiday_date,
                        'day': row['HolidayDay']
                    })
                except:
                    continue
        print(f"✅ Loaded holidays for {len(holiday_data)} states")
        return {"success": True, "states_loaded": len(holiday_data)}
    except Exception as e:
        print(f"❌ Error loading holiday data: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def load_assessment_guidelines(content):
    """Load assessment guidelines from text content"""
    global assessment_guidelines
    try:
        if isinstance(content, bytes):
            assessment_guidelines = content.decode('utf-8')
        else:
            assessment_guidelines = str(content)
        
        print(f"✅ Loaded assessment guidelines ({len(assessment_guidelines)} chars)")
        return {"success": True, "guidelines_length": len(assessment_guidelines)}
    except Exception as e:
        print(f"❌ Error loading guidelines: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def get_holidays_between_dates(state: str, start_date: datetime, end_date: datetime) -> list:
    """Get holidays for a state between two dates"""
    if state not in holiday_data:
        return []
    
    holidays = []
    for holiday in holiday_data[state]:
        if start_date <= holiday['date'] <= end_date:
            holidays.append(holiday)
    return holidays

# --- DOCUMENT DISCOVERY ---
@app.route('/get_documents', methods=['GET'])
def get_documents():
    """Fetch all unique document titles from Pinecone"""
    try:
        # Query a sample of vectors to get all unique titles
        sample_results = pinecone_index.query(
            vector=[0.0] * 768,  # Dummy vector
            top_k=1055,  # Get all records
            include_metadata=True
        )
        
        titles = set()
        for match in sample_results.get("matches", []):
            if 'metadata' in match and 'title' in match['metadata']:
                titles.add(match['metadata']['title'])
        
        sorted_titles = sorted(list(titles))
        print(f"✅ Found {len(sorted_titles)} unique documents")
        
        return jsonify({
            "documents": sorted_titles,
            "total_count": len(sorted_titles)
        }), 200
        
    except Exception as e:
        print(f"❌ Error fetching documents: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# --- CORE LOGIC ---
def perform_rag(query: str, num_neighbors: int = 5, selected_documents: list = None) -> tuple[str, list]:
    """Perform RAG search using Pinecone with optional document filtering"""
    query_embedding = embedding_model.get_embeddings([query])[0].values
    
    try:
        # Build filter if documents are specified
        query_filter = None
        if selected_documents and len(selected_documents) > 0:
            # Pinecone filter format: {"title": {"$in": ["doc1.pdf", "doc2.pdf"]}}
            query_filter = {"title": {"$in": selected_documents}}
            print(f"🔍 Filtering RAG by documents: {selected_documents}")
        
        query_results = pinecone_index.query(
            vector=query_embedding,
            top_k=num_neighbors,
            include_metadata=True,
            filter=query_filter
        )
        matches = query_results.get("matches", [])
        
        if not matches and selected_documents:
            print(f"⚠️ No matches found for selected documents. Try increasing num_neighbors or check document names.")
        
    except Exception as e:
        print(f"❌ ERROR querying Pinecone: {e}")
        traceback.print_exc()
        return "ERROR: Could not query the vector database.", []
    
    context = ""
    sources = []
    for match in matches:
        if 'metadata' in match:
            context += match['metadata'].get('text', '') + "\n---\n"
            sources.append(match['metadata'].get('title', 'Unknown'))
    
    if not context:
        context = "No relevant content found for the selected documents and query."
            
    return context, list(set(sources))

def call_openrouter(system_prompt: str, user_query: str, model: str = RAG_MODEL) -> str:
    """Call OpenRouter API for text generation with better error handling"""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://tatastrive.com",
        "X-Title": "Tata Strive RAG System"
    }
    
    data = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_query}
        ],
        "temperature": 0.3,
        "max_tokens": 3000
    }
    
    try:
        print(f"🔄 Calling OpenRouter with model: {model}")
        response = requests.post(
            f"{OPENROUTER_API_BASE}/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            print(f"❌ OpenRouter error: {response.status_code}")
            print(f"Response: {response.text}")
        
        response.raise_for_status()
        result = response.json()
        
        if "choices" not in result or len(result["choices"]) == 0:
            print(f"❌ No choices in response: {result}")
            return "Error: Invalid response from API"
        
        content = result["choices"][0]["message"]["content"]
        print(f"✅ OpenRouter response received ({len(content)} chars)")
        return content
        
    except requests.exceptions.Timeout:
        print(f"❌ OpenRouter timeout after 60 seconds")
        return "Error: API request timed out. Please try again."
    except requests.exceptions.HTTPError as e:
        print(f"❌ OpenRouter HTTP error: {e}")
        return f"Error: API returned error {e.response.status_code if hasattr(e, 'response') else 'unknown'}"
    except Exception as e:
        print(f"❌ OpenRouter error: {str(e)}")
        traceback.print_exc()
        return f"Error generating response: {str(e)}"

def translate_text(text: str, target_language: str) -> str:
    """Translate text to target language using OpenRouter"""
    if target_language == "English":
        return text
        
    system_prompt = f"Translate the following text to {target_language}. Maintain the original formatting and structure. Only provide the translation, no explanations."
    return call_openrouter(system_prompt, text, TRANSLATION_MODEL)

# --- PERSONALIZED LEARNING FUNCTIONS ---
def analyze_assessment_csv(df: pd.DataFrame) -> dict:
    """Analyze student assessment data"""
    try:
        df_latest = df.sort_values('Attempt ID').groupby(['Login ID', 'Question ID']).tail(1)
        
        student_performance = df_latest.groupby('Login ID').agg({
            'Obtained Marks': ['sum', 'count']
        }).reset_index()
        
        student_performance.columns = ['student_id', 'total_marks', 'total_questions']
        student_performance['percentage'] = (student_performance['total_marks'] / student_performance['total_questions']) * 100
        
        weak_students = student_performance[student_performance['percentage'] < 70].to_dict('records')
        
        student_details = []
        for student in weak_students:
            student_email = student['student_id']
            failed_df = df_latest[
                (df_latest['Login ID'] == student_email) & 
                (df_latest['Answer Status'] == 'Incorrect')
            ]
            
            if not failed_df.empty:
                student_details.append({
                    'email': student_email,
                    'score': f"{int(student['total_marks'])}/{int(student['total_questions'])}",
                    'percentage': round(student['percentage'], 1),
                    'failed_questions': failed_df['Question Text'].tolist()
                })
        
        question_performance = df_latest.groupby('Question Text').agg({
            'Answer Status': lambda x: ((x == 'Correct').sum() / len(x)) * 100
        }).reset_index()
        question_performance.columns = ['question', 'success_rate']
        weak_questions = question_performance[question_performance['success_rate'] < 60].to_dict('records')
        
        return {
            "total_students": len(student_performance),
            "weak_students": weak_students,
            "weak_questions": weak_questions,
            "average_score": student_performance['percentage'].mean(),
            "student_details": student_details
        }
    except Exception as e:
        print(f"Error analyzing assessment: {e}")
        traceback.print_exc()
        return {"error": str(e)}

def extract_topic_from_question(question_text: str) -> str:
    """Extract key topic from question text for RAG search"""
    clean_text = re.sub(r'<[^>]+>', '', question_text)
    clean_text = clean_text.replace('<DOUBLE_QUOTES>', '"')
    clean_text = clean_text.replace('<COMMA>', ',')
    clean_text = clean_text.replace('<br>', ' ')
    clean_text = re.sub(r'Select the correct option\..*$', '', clean_text)
    clean_text = re.sub(r'\?.*$', '', clean_text)
    
    keywords_map = {
        'hotel': 'hotel definition and types',
        'restaurant': 'restaurant and food service',
        'hospitality': 'hospitality industry basics',
        'sarai': 'traditional accommodation types in India',
        'dharamshala': 'traditional accommodation types in India',
        'front office': 'front office department and management',
        'city hotels': 'hotel classifications and types',
        'thinnai': 'traditional Indian hospitality culture',
        'independence': 'history of hotel industry in India',
        'boat houses': 'resort hotels and specialized accommodations',
        'dal lake': 'resort hotels and specialized accommodations'
    }
    
    clean_lower = clean_text.lower()
    for keyword, topic in keywords_map.items():
        if keyword in clean_lower:
            return f"Front Desk Associate {topic}"
    
    return "Front Desk Associate hospitality basics"

def clean_question_text(question_text: str) -> str:
    """Clean question text for display"""
    clean = re.sub(r'<[^>]+>', '', question_text)
    clean = clean.replace('<DOUBLE_QUOTES>', '"')
    clean = clean.replace('<COMMA>', ',')
    clean = clean.replace('<br>', ' ')
    if len(clean) > 100:
        clean = clean[:100] + "..."
    return clean

def generate_personalized_content_for_student(student_email: str, failed_questions: list) -> str:
    """Generate personalized study content for a student"""
    topics = [extract_topic_from_question(q) for q in failed_questions]
    weak_topics = ", ".join(set(topics))
    
    all_context = ""
    for topic in set(topics):
        context, _ = perform_rag(topic, num_neighbors=3)
        all_context += context + "\n\n"
    
    prompt = f"""You are an expert educator creating personalized remedial content for a Tata Strive Front Desk Associate student.

Student: {student_email}
Topics they struggled with: {weak_topics}

Questions they got wrong:
{chr(10).join([f"- {q}" for q in failed_questions])}

Context from knowledge base:
{all_context}

Create a comprehensive study guide that:
1. Explains each concept in simple, clear language
2. Addresses common misconceptions
3. Provides practical examples relevant to Front Desk Associate work
4. Includes memory aids and tips
5. Adds 5 practice questions with detailed explanations
6. Uses encouraging, supportive tone

Focus on helping them understand these specific topics better."""
    
    content = call_openrouter(prompt, "Generate personalized study guide")
    return content

def send_email_with_pdf(to_email: str, subject: str, body: str, pdf_content: bytes, pdf_name: str) -> bool:
    """Send HTML email with PDF attachment using Gmail SMTP"""
    try:
        msg = MIMEMultipart('alternative')
        msg['From'] = f"{EMAIL_SENDER_NAME} <{EMAIL_SENDER}>"
        msg['To'] = to_email
        msg['Subject'] = subject
        
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #2c3e50;">Hi Student,</h2>
                
                <p>You recently completed the <strong>Front Desk Associate</strong> assessment.</p>
                
                {body}
                
                <div style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 12px;">
                    <p>This is an automated message from Tata Strive Learning Platform.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(body, 'plain'))
        msg.attach(MIMEText(html_body, 'html'))
        
        pdf_attachment = MIMEApplication(pdf_content, _subtype="pdf")
        pdf_attachment.add_header('Content-Disposition', 'attachment', filename=pdf_name)
        msg.attach(pdf_attachment)
        
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.send_message(msg)
        
        print(f"✅ Email sent successfully to {to_email}")
        return True
        
    except Exception as e:
        print(f"❌ Failed to send email to {to_email}: {str(e)}")
        traceback.print_exc()
        return False

# --- PROMPTS ---
ASSESSMENT_PROMPT = """You are an expert educational assessment creator for Tata Strive programs.

GUIDELINES TO FOLLOW:
{guidelines}

Context from knowledge base:
{context}

Create an assessment that:
1. Follows ALL guidelines strictly (MCSS/MCMS format, 4 options, no "All of the above")
2. Uses Bloom's taxonomy levels 1-2 (Knowledge and Comprehension)
3. Tests practical understanding
4. Provides clear marking criteria
5. Is culturally sensitive and uses inclusive language

Format the output clearly with numbered questions and marking schemes."""

LESSON_PLAN_PROMPT = """You are an expert curriculum designer for Tata Strive programs.

Course Information:
{course_info}

Holidays to Consider:
{holidays}

Context from knowledge base:
{context}

Create a detailed lesson plan that:
1. Spans the exact course duration ({duration} hours)
2. EXCLUDES all holidays listed above
3. EXCLUDES weekends (Saturday/Sunday)
4. Includes:
   - Learning objectives
   - Daily schedule with time allocation
   - Module-wise breakdown
   - Assessment strategies
   - Materials needed
   - Key takeaways

Make it practical, realistic, and aligned with Tata Strive's teaching methodology.
Ensure the timeline accounts for holidays and weekends."""

# --- DOCUMENT GENERATION ---
def generate_docx(content: str, title: str) -> io.BytesIO:
    """Generate a DOCX document from content"""
    doc = Document()
    doc.add_heading(title, 0)
    
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def generate_pdf(content: str, title: str) -> io.BytesIO:
    """Generate a PDF document from content"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(10)
    
    pdf.set_font("helvetica", size=11)
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            clean_line = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, clean_line)
            pdf.ln(2)
    
    pdf_buffer = io.BytesIO()
    pdf_content = pdf.output(dest='S').encode('latin-1')
    pdf_buffer.write(pdf_content)
    pdf_buffer.seek(0)
    return pdf_buffer

# --- ENDPOINTS ---
@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy", 
        "service": "Tata Strive RAG API",
        "courses_loaded": len(course_data),
        "states_with_holidays": len(holiday_data),
        "guidelines_loaded": len(assessment_guidelines) > 0
    }), 200

@app.route('/detect_location', methods=['POST'])
def detect_location():
    """Detect location from IP and suggest language"""
    try:
        data = request.json
        ip_address = data.get('ip', request.remote_addr)
        
        location = detect_location_from_ip(ip_address)
        suggested_lang = get_suggested_language(location['state'])
        
        return jsonify({
            "location": location,
            "suggested_language": suggested_lang
        }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload/course_data', methods=['POST'])
def upload_course_data():
    """Upload course duration data"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        result = load_course_data(file)
        return jsonify(result), 200 if result['success'] else 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload/holidays', methods=['POST'])
def upload_holidays():
    """Upload holiday data"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        result = load_holiday_data(file)
        return jsonify(result), 200 if result['success'] else 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload/guidelines', methods=['POST'])
def upload_guidelines():
    """Upload assessment guidelines"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        result = load_assessment_guidelines(file.read())
        return jsonify(result), 200 if result['success'] else 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/create/assessment', methods=['POST'])
def create_assessment():
    """Create an assessment based on a topic"""
    try:
        data = request.json
        topic = data.get('query', '')
        requirements = data.get('requirements', '')
        language = data.get('language', 'English')
        format_type = data.get('output_format', 'json')
        selected_documents = data.get('selected_documents', [])
        
        if not topic:
            return jsonify({"error": "Query/topic is required"}), 400
        
        if not selected_documents or len(selected_documents) == 0:
            return jsonify({"error": "Please select at least one source document"}), 400
        
        context, sources = perform_rag(topic, num_neighbors=10, selected_documents=selected_documents)
        full_query = f"Topic: {topic}\nRequirements: {requirements}"
        
        prompt = ASSESSMENT_PROMPT.format(
            guidelines=assessment_guidelines if assessment_guidelines else "Standard MCQ format with 4 options.",
            context=context
        )
        
        assessment = call_openrouter(prompt, full_query)
        translated_assessment = translate_text(assessment, language)
        
        if format_type == 'docx':
            doc_buffer = generate_docx(translated_assessment, f"Assessment: {topic}")
            return send_file(doc_buffer, as_attachment=True, 
                           download_name=f"assessment_{topic.replace(' ', '_')}.docx",
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif format_type == 'pdf':
            pdf_buffer = generate_pdf(translated_assessment, f"Assessment: {topic}")
            return send_file(pdf_buffer, as_attachment=True,
                           download_name=f"assessment_{topic.replace(' ', '_')}.pdf",
                           mimetype='application/pdf')
        
        return jsonify({
            "english_answer": assessment,
            "translated_answer": translated_assessment,
            "language": language,
            "sources": sources
        }), 200
        
    except Exception as e:
        print(f"Error in create_assessment: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/create/lesson_plan', methods=['POST'])
def create_lesson_plan():
    """Create a lesson plan based on a topic"""
    try:
        data = request.json
        topic = data.get('query', '')
        course_name = data.get('course_name', '')
        state = data.get('state', 'Corporate')
        start_date_str = data.get('start_date', '')
        language = data.get('language', 'English')
        format_type = data.get('output_format', 'json')
        selected_documents = data.get('selected_documents', [])
        
        if not topic:
            return jsonify({"error": "Topic is required"}), 400
        
        if not selected_documents or len(selected_documents) == 0:
            return jsonify({"error": "Please select at least one source document"}), 400
        
        course_info_str = "Duration not specified"
        duration_hours = 0
        if course_name and course_name in course_data:
            course_info = course_data[course_name]
            duration_hours = course_info['duration_hours']
            course_info_str = f"Course: {course_name}\nDuration: {duration_hours} hours\nTheory Hours: {course_info['theory_hours']}\nEligibility: {course_info['eligibility']}"
        
        holidays_str = "No holidays data available"
        if state and start_date_str and state in holiday_data:
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
                days_needed = int(duration_hours / 7) if duration_hours > 0 else 30
                end_date = start_date + timedelta(days=days_needed * 2)
                
                holidays = get_holidays_between_dates(state, start_date, end_date)
                if holidays:
                    holidays_str = "Holidays to EXCLUDE:\n" + "\n".join([
                        f"- {h['name']}: {h['date'].strftime('%d %B %Y')} ({h['day']})"
                        for h in holidays
                    ])
            except:
                pass
        
        context, sources = perform_rag(topic, num_neighbors=10, selected_documents=selected_documents)
        
        full_query = f"Topic: {topic}\nCreate lesson plan considering all holidays and weekends."
        prompt = LESSON_PLAN_PROMPT.format(
            course_info=course_info_str,
            duration=f"{duration_hours}" if duration_hours > 0 else "unspecified",
            holidays=holidays_str,
            context=context
        )
        
        lesson_plan = call_openrouter(prompt, full_query)
        translated_lesson_plan = translate_text(lesson_plan, language)
        
        if format_type == 'docx':
            doc_buffer = generate_docx(translated_lesson_plan, f"Lesson Plan: {topic}")
            return send_file(doc_buffer, as_attachment=True,
                           download_name=f"lesson_plan_{topic.replace(' ', '_')}.docx",
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif format_type == 'pdf':
            pdf_buffer = generate_pdf(translated_lesson_plan, f"Lesson Plan: {topic}")
            return send_file(pdf_buffer, as_attachment=True,
                           download_name=f"lesson_plan_{topic.replace(' ', '_')}.pdf",
                           mimetype='application/pdf')
        
        return jsonify({
            "english_answer": lesson_plan,
            "translated_answer": translated_lesson_plan,
            "language": language,
            "sources": sources,
            "holidays_considered": holidays_str
        }), 200
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- PERSONALIZED LEARNING ENDPOINT ---
@app.route('/process/assessment_and_email', methods=['POST'])
def process_assessment_and_email():
    """Process assessment CSV and automatically email personalized content to weak students"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
        else:
            return jsonify({"error": "Invalid file format. Upload CSV or Excel"}), 400
        
        print("📊 Analyzing assessment data...")
        analysis = analyze_assessment_csv(df)
        
        if 'error' in analysis:
            return jsonify({"error": analysis['error']}), 500
        
        email_results = []
        student_details = analysis.get('student_details', [])
        
        print(f"📧 Processing {len(student_details)} students who need support...")
        
        for student in student_details:
            try:
                student_email = student['email']
                failed_questions = student['failed_questions']
                score = student['score']
                percentage = student['percentage']
                
                print(f"Generating content for {student_email}...")
                
                content = generate_personalized_content_for_student(student_email, failed_questions)
                
                pdf_buffer = generate_pdf(content, f"Personalized Study Guide - {student_email}")
                pdf_bytes = pdf_buffer.getvalue()
                
                cleaned_questions = [clean_question_text(q) for q in failed_questions[:5]]
                
                body_plain = f"""Hi Student,\n\nYou recently completed the Front Desk Associate assessment.\n\nYOUR RESULTS:\nScore: {score} ({percentage}%)\nStatus: Needs Improvement\n\n..."""

                body_html = f"""
                <div style="background-color: #f8f9fa; padding: 20px; border-radius: 5px; margin: 20px 0;">
                    <h3 style="color: #2c3e50; margin-top: 0;">📊 YOUR RESULTS:</h3>
                    <p style="font-size: 18px;"><strong>Score:</strong> {score} ({percentage}%)</p>
                    <p style="color: #e74c3c;"><strong>Status:</strong> Needs Improvement</p>
                </div>
                <p>We've analyzed your performance and created a personalized study guide to help you master the concepts you found challenging.</p>
                <div style="background-color: #fff3cd; padding: 15px; border-left: 4px solid #ffc107; margin: 20px 0;">
                    <h4 style="color: #856404; margin-top: 0;">⚠️ TOPICS YOU STRUGGLED WITH:</h4>
                    <ul style="color: #856404;">
                        {chr(10).join([f"<li>{q}</li>" for q in cleaned_questions])}
                    </ul>
                </div>
                <div style="background-color: #d4edda; padding: 15px; border-radius: 5px; margin: 20px 0;">
                    <h4 style="color: #155724; margin-top: 0;">📎 ATTACHED: Your_Personalized_Study_Guide.pdf</h4>
                    <p style="color: #155724; margin-bottom: 10px;"><strong>This guide includes:</strong></p>
                    <ul style="color: #155724;">
                        <li>Clear explanations of each topic</li>
                        <li>Practical examples for Front Desk work</li>
                        <li>Memory tips and tricks</li>
                        <li>Practice questions with answers</li>
                    </ul>
                </div>
                <p style="background-color: #e3f2fd; padding: 10px; border-left: 4px solid #2196f3; margin: 20px 0;">
                    <strong>💡 TIP:</strong> Review this guide before your next attempt!
                </p>
                <p style="margin-top: 30px;">Best regards,<br><strong>{EMAIL_SENDER_NAME}</strong></p>
                """
                
                subject = "📚 Your Personalized Study Guide - Front Desk Associate"
                
                success = send_email_with_pdf(
                    to_email=student_email,
                    subject=subject,
                    body=body_html,
                    pdf_content=pdf_bytes,
                    pdf_name=f"Study_Guide_{student_email.split('@')[0]}.pdf"
                )
                
                email_results.append({
                    "email": student_email,
                    "status": "✅ Sent" if success else "❌ Failed",
                    "score": score,
                    "percentage": percentage
                })
                
            except Exception as e:
                print(f"Error processing student {student['email']}: {str(e)}")
                email_results.append({
                    "email": student['email'],
                    "status": f"❌ Error: {str(e)}",
                    "score": student.get('score', 'N/A'),
                    "percentage": student.get('percentage', 'N/A')
                })
        
        return jsonify({
            "total_students": analysis['total_students'],
            "average_score": round(analysis['average_score'], 1),
            "emails_sent": len([r for r in email_results if "✅" in r['status']]),
            "email_results": email_results,
            "weak_questions": analysis['weak_questions']
        }), 200
        
    except Exception as e:
        print(f"Error in process_assessment_and_email: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/search', methods=['POST'])
def search():
    """Direct search endpoint for testing"""
    try:
        data = request.json
        query = data.get('query', '')
        selected_documents = data.get('selected_documents', None)
        
        if not query:
            return jsonify({"error": "Query is required"}), 400
        
        context, sources = perform_rag(query, num_neighbors=5, selected_documents=selected_documents)
        
        return jsonify({
            "context": context,
            "sources": sources
        }), 200
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("Starting Tata Strive RAG API Server...")
    print("=" * 50)
    print("🚀 Tata Strive RAG API Server")
    print("=" * 50)
    print(f"📊 Pinecone Index: {PINECONE_INDEX_NAME}")
    print(f"🤖 LLM Model: {RAG_MODEL}")
    print(f"🌐 Translation Model: {TRANSLATION_MODEL}")
    print(f"📧 Email Sender: {EMAIL_SENDER}")
    print(f"🔧 Server running on: http://0.0.0.0:8081")
    print("=" * 50)
    print("\nAvailable Endpoints:")
    print("  GET  /health")
    print("  GET  /get_documents")
    print("  POST /detect_location")
    print("  POST /upload/course_data")
    print("  POST /upload/holidays")
    print("  POST /upload/guidelines")
    print("  POST /create/assessment")
    print("  POST /create/lesson_plan")
    print("  POST /process/assessment_and_email")
    print("  POST /search")
    print("=" * 50)
    print("\n✅ Server is ready to accept requests!\n")
    
    app.run(host='0.0.0.0', port=8081, debug=True)