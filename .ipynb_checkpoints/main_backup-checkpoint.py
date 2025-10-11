# main.py - Final Development Version with Local FAISS Search

# --- Imports ---
import re
import os
import io
import json
import numpy as np
import faiss
import traceback
from dotenv import load_dotenv
from flask import Flask, request, jsonify, send_file
from google.cloud import storage
import vertexai
from vertexai.language_models import TextEmbeddingModel
import requests
from docx import Document
from fpdf2 import FPDF
from fpdf2.enums import XPos, YPos

# --- Initialization & Configuration ---
load_dotenv()

# GCP & Model Config
PROJECT_ID = os.environ.get("GOOGLE_CLOUD_PROJECT")
LOCATION = "asia-south1"
BUCKET_NAME = "rag-source-documents"
METADATA_FILE_PATH = "vector-search-inputs/metadata_lookup.json"
VECTOR_INPUT_FILE_PATH = "vector-search-inputs/vector_search_input.json"
LOCAL_INDEX_FILE = "local_app_index.faiss"

# OpenRouter & Model Config
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"
EMBEDDING_MODEL = "text-embedding-005"
RAG_MODEL = "tngtech/deepseek-r1t2-chimera:free"
TRANSLATION_MODEL = "openrouter/sonoma-sky-alpha"

# --- Client Initialization ---
vertexai.init(project=PROJECT_ID, location=LOCATION)
storage_client = storage.Client()
embedding_model = TextEmbeddingModel.from_pretrained(EMBEDDING_MODEL)
app = Flask(__name__)

# --- Data Loading ---

def load_json_from_gcs(bucket_name, file_path):
    """Downloads and loads a JSON or JSONL file from GCS."""
    print(f"Loading {file_path} from gs://{bucket_name}...")
    try:
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(file_path)
        content = blob.download_as_string()
        if "vector_search_input" in file_path:
            return [json.loads(line) for line in content.decode('utf-8').splitlines()]
        else:
            return json.loads(content)
    except Exception as e:
        print(f"❌ FATAL: Could not load file from GCS. Error: {e}")
        return None

metadata_lookup = load_json_from_gcs(BUCKET_NAME, METADATA_FILE_PATH)
id_lookup_data = load_json_from_gcs(BUCKET_NAME, VECTOR_INPUT_FILE_PATH)
id_lookup = [item['id'] for item in id_lookup_data] if id_lookup_data else []

print(f"Loading local FAISS index from {LOCAL_INDEX_FILE}...")
try:
    local_index = faiss.read_index(LOCAL_INDEX_FILE)
    print("✅ Local FAISS index loaded successfully.")
except Exception as e:
    print(f"❌ FATAL: Could not load local FAISS index. Did you run build_local_index.py? Error: {e}")
    local_index = None

# --- Prompt Engineering ---
PROMPT_TEMPLATES = {
    "assessment": """You are an expert quiz designer for Tata Strive. Your user is in {location}. Based ONLY on the provided context, create a multiple-choice quiz with 5 questions. Make the questions culturally and geographically relevant to {location}. Each question must have 4 options, with one correct answer clearly marked with "(Correct Answer)".
CONTEXT: {context}
USER REQUEST: {user_query}
QUIZ:""",
    "lesson_plan": """You are an instructional designer for Tata Strive creating a lesson plan for facilitators in {location}. Using ONLY the provided context, create a structured 30-minute lesson plan. The plan must include a clear learning objective, a list of materials, an engaging activity that would resonate with learners in {location}, and a method for checking understanding.
CONTEXT: {context}
USER REQUEST: {user_query}
LESSON PLAN:""",
    "content_generator": """You are a creative educational content writer for Tata Strive. Your target audience is in {location}. Using the provided context as a base, generate a supplementary article that expands on the user's request. Make it engaging and use examples relevant to {location}.
CONTEXT: {context}
USER REQUEST: {user_query}
ARTICLE:"""
}

# --- Document Generation Tools ---

def create_docx(content: str, title: str) -> io.BytesIO:
    """Generates a .docx file, correctly handling **bold** markdown."""
    document = Document()
    document.add_heading(title, level=1)
    content = content.strip()
    for paragraph_text in content.split('\n'):
        if not paragraph_text.strip():
            continue
        p = document.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', paragraph_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
    mem_file = io.BytesIO()
    document.save(mem_file)
    mem_file.seek(0)
    return mem_file

def create_pdf(content: str, title: str) -> io.BytesIO:
    """Generates a professionally formatted PDF from markdown-like text content."""
    pdf = FPDF()
    pdf.add_page()
    dejavu_font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
    if os.path.exists(dejavu_font_path):
        pdf.add_font('DejaVu', '', dejavu_font_path)
        pdf.set_font('DejaVu', '', 12)
    else:
        print("⚠️ WARNING: DejaVuSans.ttf not found. Using Helvetica.")
        pdf.set_font("Helvetica", size=12)
    pdf.set_font(style="B", size=16)
    pdf.multi_cell(0, 10, text=title, align='C')
    pdf.ln(10)
    pdf.set_font(size=12)
    lines = content.split('\n')
    QUESTION_BLOCK_HEIGHT = 45
    for line in lines:
        line = line.strip()
        if not line or line.isdigit():
            continue
        if re.match(r'^\d+\.', line):
            if pdf.will_page_break(QUESTION_BLOCK_HEIGHT):
                pdf.add_page()
            pdf.ln(5)
            pdf.set_font(style="B")
            pdf.multi_cell(0, 7, text=line)
            pdf.set_font(style="")
        elif line.startswith(("*Why?*", "*Correct Answer:")):
            pdf.ln(2)
            pdf.set_font(style="I")
            pdf.multi_cell(0, 7, text=line)
            pdf.set_font(style="")
        elif re.match(r'^[A-D]\)', line) or "(Correct Answer)" in line:
            is_correct = "(Correct Answer)" in line
            if is_correct:
                pdf.set_font(style="B")
            pdf.set_x(15)
            pdf.multi_cell(0, 7, text=line)
            if is_correct:
                pdf.set_font(style="")
        else:
            pdf.multi_cell(0, 7, text=line)
    mem_file = io.BytesIO(pdf.output())
    mem_file.seek(0)
    return mem_file

# --- Core Logic Functions ---

def perform_rag(query: str, num_neighbors: int = 5) -> tuple[str, list]:
    """Performs RAG using the local FAISS index."""
    if not local_index:
        return "ERROR: Local index is not loaded.", []
    query_embedding = embedding_model.get_embeddings([query])[0].values
    query_vector = np.array([query_embedding]).astype('float32')
    distances, indices = local_index.search(query_vector, num_neighbors)
    context, sources = "", []
    for i in indices[0]:
        if i != -1:
            chunk_id = id_lookup[i]
            if chunk_id in metadata_lookup:
                chunk_data = metadata_lookup[chunk_id]
                context += chunk_data.get('content', '') + "\n---\n"
                sources.append(chunk_data.get('title', 'Unknown'))
    return context, list(set(sources))

def call_openrouter(prompt: str, model: str) -> str:
    """Calls the OpenRouter API for generation."""
    response = requests.post(
        url=f"{OPENROUTER_API_BASE}/chat/completions",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}"},
        json={"model": model, "messages": [{"role": "user", "content": prompt}]}
    )
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content']

def translate_text(text: str, language: str) -> str:
    """Translates text using the specified OpenRouter model."""
    if not text or language.lower() in ['en', 'english']:
        return text
    prompt = f"Translate the following English text into {language}. Provide only the direct translation.\n\n{text}"
    return call_openrouter(prompt, TRANSLATION_MODEL)

# --- API Handler & Endpoints ---

def handle_creation_request(feature_type: str):
    """Generic handler for all creation requests."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON request"}), 400
        
        user_query = data.get('query')
        if not user_query:
            return jsonify({"error": "Request must include 'query'"}), 400

        target_language = data.get('language', 'English')
        location = data.get('location', 'India')
        output_format = data.get('output_format', 'json').lower()
        
        context, sources = perform_rag(user_query)
        if not context.strip():
            return jsonify({"error": "Could not find any relevant context for the query."}), 404

        prompt = PROMPT_TEMPLATES[feature_type].format(context=context, user_query=user_query, location=location)
        
        english_answer = call_openrouter(prompt, RAG_MODEL)
        final_content = translate_text(english_answer, target_language)

        if output_format == 'json':
            return jsonify({
                "english_answer": english_answer, "translated_answer": final_content,
                "language": target_language, "sources": sources
            })
        
        title = f"Tata Strive - {feature_type.replace('_', ' ').title()}"
        if output_format == 'docx':
            file_stream = create_docx(final_content, title)
            return send_file(file_stream, as_attachment=True, download_name=f'{feature_type}.docx',
                             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        if output_format == 'pdf':
            file_stream = create_pdf(final_content, title)
            return send_file(file_stream, as_attachment=True, download_name=f'{feature_type}.pdf', mimetype='application/pdf')

        return jsonify({"error": "Unsupported output_format. Use 'json', 'docx', or 'pdf'."}), 400

    except Exception as e:
        print(f"❌ An unexpected error occurred: {e}")
        traceback.print_exc()
        return jsonify({"error": "An internal server error occurred.", "details": str(e)}), 500

@app.route('/create/assessment', methods=['POST'])
def create_assessment_endpoint():
    return handle_creation_request('assessment')

@app.route('/create/lesson_plan', methods=['POST'])
def create_lesson_plan_endpoint():
    return handle_creation_request('lesson_plan')

@app.route('/create/content', methods=['POST'])
def create_content_endpoint():
    return handle_creation_request('content_generator')

# --- Main Execution ---
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=8081)