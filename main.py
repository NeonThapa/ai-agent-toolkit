import os
import io
import json
import re
import html
import traceback
from pathlib import Path
import pandas as pd
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime, timedelta
from dotenv import load_dotenv
from pinecone import Pinecone
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from google.cloud import storage
import vertexai
from vertexai.language_models import TextEmbeddingModel
import requests
from docx import Document
from docx.shared import Pt
from markdown_it import MarkdownIt

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
    WEASYPRINT_IMPORT_ERROR = None
except Exception as _weasy_err:
    HTML = CSS = None
    WEASYPRINT_AVAILABLE = False
    WEASYPRINT_IMPORT_ERROR = _weasy_err

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# --- CONFIGURATION ---
load_dotenv()
PROJECT_ID = os.environ.get("GOOGLE_CLOUD_PROJECT")
LOCATION = "asia-south1"
BUCKET_NAME = "rag-source-documents"
METADATA_FILE_PATH = "vector-search-inputs/metadata_lookup.json"
PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY")
PINECONE_INDEX_NAME = "tata-strive-rag"
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"
EMBEDDING_MODEL = "text-embedding-005"
RAG_MODEL = os.environ.get("RAG_MODEL", "tngtech/deepseek-r1t2-chimera:free")
TRANSLATION_MODEL = os.environ.get("TRANSLATION_MODEL", "z-ai/glm-4.5-air:free")

# Email Configuration
EMAIL_SENDER = os.environ.get("EMAIL_SENDER_ADDRESS")
EMAIL_PASSWORD = os.environ.get("EMAIL_SENDER_PASSWORD")
EMAIL_SENDER_NAME = os.environ.get("EMAIL_SENDER_NAME", "Tata Strive Learning Team")
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587

# --- IN-MEMORY DATA STORAGE ---
course_data = {}  # Will store course duration info
holiday_data = {}  # Will store holidays by state
assessment_guidelines = ""  # Will store assessment guidelines

# --- CLIENT INITIALIZATION ---
vertexai.init(project=PROJECT_ID, location=LOCATION)
storage_client = storage.Client()
embedding_model = TextEmbeddingModel.from_pretrained(EMBEDDING_MODEL)
app = Flask(__name__)

raw_origins = os.environ.get("FRONTEND_ORIGINS", "*")
if raw_origins == "*":
    allowed_origins = "*"
else:
    allowed_origins = [origin.strip() for origin in raw_origins.split(",") if origin.strip()]

CORS(
    app,
    resources={r"/*": {"origins": allowed_origins}},
    supports_credentials=False,
    allow_headers=["*"],
    methods=["GET", "POST", "OPTIONS"]
)

markdown_renderer = MarkdownIt("commonmark").enable("table").enable("strikethrough").enable("linkify")

PDF_BASE_STYLE_CSS = """
@page {
    size: A4;
    margin: 1in 0.9in 1in 0.9in;
}
body {
    font-family: 'Calibri', 'Segoe UI', 'DejaVu Sans', Arial, sans-serif;
    color: #1f2933;
    line-height: 1.65;
    font-size: 12pt;
}
h1 {
    font-size: 26pt;
    margin-bottom: 0.3em;
    color: #005b99;
    letter-spacing: -0.02em;
}
h2 {
    font-size: 20pt;
    margin-top: 1.2em;
    margin-bottom: 0.25em;
    color: #0f172a;
}
h3 {
    font-size: 16pt;
    margin-top: 1em;
    margin-bottom: 0.2em;
    color: #0f172a;
}
h4, h5, h6 {
    margin-top: 0.7em;
    margin-bottom: 0.15em;
    color: #0f172a;
}
p {
    margin: 0.45em 0;
}
strong {
    color: #0f172a;
}
em {
    color: #1d4ed8;
}
blockquote {
    border-left: 4px solid rgba(0, 176, 255, 0.55);
    padding-left: 14px;
    margin: 1em 0;
    color: #334155;
    font-style: italic;
    background: rgba(0, 176, 255, 0.06);
}
ul, ol {
    margin: 0.4em 0 0.6em 1.2em;
}
li {
    margin: 0.25em 0;
}
table {
    width: 100%;
    border-collapse: collapse;
    margin: 1.1em 0;
    font-size: 11pt;
}
thead tr {
    background: rgba(0, 176, 255, 0.12);
}
th, td {
    border: 1px solid #dbeafe;
    padding: 8px 10px;
    vertical-align: top;
}
tbody tr:nth-child(even) {
    background: #f4f9ff;
}
code {
    font-family: 'Fira Code', 'Consolas', 'Courier New', monospace;
    background: rgba(15, 23, 42, 0.08);
    padding: 2px 5px;
    border-radius: 4px;
}
.title-header {
    border-bottom: 3px solid rgba(0, 176, 255, 0.6);
    padding-bottom: 0.6em;
    margin-bottom: 1.4em;
}
.subtitle {
    font-size: 12pt;
    color: #475569;
    margin: 0;
}
.content {
    margin-top: 0.8em;
}
"""

FONT_DIR = Path(__file__).resolve().parent / "assets" / "fonts"
PDF_FONT_DEFAULT = "NotoSans"
PDF_FONT_BOLD = "NotoSans-Bold"
PDF_FONT_DEVANAGARI = "NotoSansDevanagari"
PDF_FONT_DEVANAGARI_BOLD = "NotoSansDevanagari-Bold"
PDF_FONT_BENGALI = "NotoSansBengali"
PDF_FONT_BENGALI_BOLD = "NotoSansBengali-Bold"
PDF_FONT_TAMIL = "NotoSansTamil"
PDF_FONT_TAMIL_BOLD = "NotoSansTamil-Bold"
PDF_FONT_TELUGU = "NotoSansTelugu"
PDF_FONT_TELUGU_BOLD = "NotoSansTelugu-Bold"
PDF_FONT_GUJARATI = "NotoSansGujarati"
PDF_FONT_GUJARATI_BOLD = "NotoSansGujarati-Bold"
PDF_FONT_KANNADA = "NotoSansKannada"
PDF_FONT_KANNADA_BOLD = "NotoSansKannada-Bold"

_SCRIPT_PATTERNS = {
    "bengali": re.compile(r"[\u0980-\u09FF]"),
    "devanagari": re.compile(r"[\u0900-\u097F]"),
    "tamil": re.compile(r"[\u0B80-\u0BFF]"),
    "telugu": re.compile(r"[\u0C00-\u0C7F]"),
    "gujarati": re.compile(r"[\u0A80-\u0AFF]"),
    "kannada": re.compile(r"[\u0C80-\u0CFF]"),
}
_PDF_FONTS_REGISTERED = False
_AVAILABLE_PDF_FONTS: set[str] = set()
_PDF_STYLE_CACHE: dict[str, ParagraphStyle] = {}

WEASYPRINT_WARNING_EMITTED = False


def _register_pdf_fonts() -> None:
    global _PDF_FONTS_REGISTERED
    if _PDF_FONTS_REGISTERED:
        return

    font_files = {
        PDF_FONT_DEFAULT: "NotoSans-Regular.ttf",
        PDF_FONT_BOLD: "NotoSans-Bold.ttf",
        PDF_FONT_DEVANAGARI: "NotoSansDevanagari-Regular.ttf",
        PDF_FONT_DEVANAGARI_BOLD: "NotoSansDevanagari-Bold.ttf",
        PDF_FONT_BENGALI: "NotoSansBengali-Regular.ttf",
        PDF_FONT_BENGALI_BOLD: "NotoSansBengali-Bold.ttf",
        PDF_FONT_TAMIL: "NotoSansTamil-Regular.ttf",
        PDF_FONT_TAMIL_BOLD: "NotoSansTamil-Bold.ttf",
        PDF_FONT_TELUGU: "NotoSansTelugu-Regular.ttf",
        PDF_FONT_TELUGU_BOLD: "NotoSansTelugu-Bold.ttf",
        PDF_FONT_GUJARATI: "NotoSansGujarati-Regular.ttf",
        PDF_FONT_GUJARATI_BOLD: "NotoSansGujarati-Bold.ttf",
        PDF_FONT_KANNADA: "NotoSansKannada-Regular.ttf",
        PDF_FONT_KANNADA_BOLD: "NotoSansKannada-Bold.ttf",
    }

    for font_name, filename in font_files.items():
        font_path = FONT_DIR / filename
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                _AVAILABLE_PDF_FONTS.add(font_name)
            except Exception as font_err:
                print(f"⚠️ Unable to register font '{font_name}': {font_err}")
        else:
            print(f"⚠️ Font file not found for '{font_name}' at {font_path}")

    if PDF_FONT_DEFAULT in _AVAILABLE_PDF_FONTS and PDF_FONT_BOLD in _AVAILABLE_PDF_FONTS:
        pdfmetrics.registerFontFamily(
            "TataSans",
            normal=PDF_FONT_DEFAULT,
            bold=PDF_FONT_BOLD,
            italic=PDF_FONT_DEFAULT,
            boldItalic=PDF_FONT_BOLD,
        )
    else:
        print("⚠️ Core Noto Sans fonts missing; PDF output may not render as expected.")

    _PDF_FONTS_REGISTERED = True


def _detect_script(text: str) -> str:
    checks = _SCRIPT_PATTERNS
    for key, pattern in checks.items():
        if pattern.search(text):
            return key
    return "latin"


def _select_pdf_font(text: str, bold: bool = False) -> str:
    script = _detect_script(text)
    if script == "bengali":
        font_name = PDF_FONT_BENGALI_BOLD if bold else PDF_FONT_BENGALI
    elif script == "devanagari":
        font_name = PDF_FONT_DEVANAGARI_BOLD if bold else PDF_FONT_DEVANAGARI
    elif script == "tamil":
        font_name = PDF_FONT_TAMIL_BOLD if bold else PDF_FONT_TAMIL
    elif script == "telugu":
        font_name = PDF_FONT_TELUGU_BOLD if bold else PDF_FONT_TELUGU
    elif script == "gujarati":
        font_name = PDF_FONT_GUJARATI_BOLD if bold else PDF_FONT_GUJARATI
    elif script == "kannada":
        font_name = PDF_FONT_KANNADA_BOLD if bold else PDF_FONT_KANNADA
    else:
        font_name = PDF_FONT_BOLD if bold else PDF_FONT_DEFAULT
    if font_name not in _AVAILABLE_PDF_FONTS:
        fallback = PDF_FONT_BOLD if bold else PDF_FONT_DEFAULT
        return fallback if fallback in _AVAILABLE_PDF_FONTS else font_name
    return font_name


def _style_for_text(base_style: ParagraphStyle, text: str, force_bold: bool | None = None) -> ParagraphStyle:
    normalized = _ensure_text(text)
    font_is_bold = force_bold if force_bold is not None else base_style.fontName.endswith("-Bold")
    font_name = _select_pdf_font(normalized, font_is_bold)
    if base_style.fontName == font_name:
        return base_style

    cache_key = f"{base_style.name}|{font_name}"
    if cache_key not in _PDF_STYLE_CACHE:
        _PDF_STYLE_CACHE[cache_key] = ParagraphStyle(
            cache_key,
            parent=base_style,
            fontName=font_name,
        )
    return _PDF_STYLE_CACHE[cache_key]

def _ensure_text(value) -> str:
    """Coerce arbitrary content into safe unicode text."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, bytes):
        try:
            return value.decode("utf-8", errors="ignore")
        except Exception:
            return value.decode("latin-1", errors="ignore")
    return str(value)


def _render_inline_html(text) -> str:
    snippet = markdown_renderer.render(_ensure_text(text)).strip()
    if snippet.startswith("<p>") and snippet.endswith("</p>"):
        snippet = snippet[3:-4]
    return snippet or ""


def _build_list_flowable(items: list[str], list_kind: str, item_style: ParagraphStyle) -> list:
    """Render list items as indented paragraphs without auto-numbering artifacts."""
    flowables: list = []
    for item in items:
        text = _ensure_text(item).strip()
        if not text:
            continue
        style = _style_for_text(item_style, text)
        flowables.append(Paragraph(_render_inline_html(text), style))
    if flowables:
        flowables.append(Spacer(1, 4))
    return flowables


def _build_table_flowable(block: dict, header_style: ParagraphStyle, body_style: ParagraphStyle) -> Table:
    header = [_ensure_text(cell) for cell in (block.get('header', []) or [])]
    rows = [
        [_ensure_text(cell) for cell in (_row or [])]
        for _row in (block.get('rows', []) or [])
    ]

    table_text_rows = []
    if header:
        table_text_rows.append(header)
    table_text_rows.extend(rows)

    if not table_text_rows:
        table_text_rows = [["No data available"]]

    max_cols = max(len(r) for r in table_text_rows) if table_text_rows else 1
    normalized_rows = []
    for idx, row in enumerate(table_text_rows):
        padded = row + [""] * (max_cols - len(row))
        is_header_row = header and idx == 0
        base_style = header_style if is_header_row else body_style
        normalized_row = []
        for cell in padded:
            cell_style = _style_for_text(base_style, cell, force_bold=is_header_row)
            normalized_row.append(Paragraph(_render_inline_html(cell), cell_style))
        normalized_rows.append(normalized_row)

    table = Table(normalized_rows, hAlign='LEFT')
    table_styles = [
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('INNERGRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#bfdbfe')),
        ('BOX', (0, 0), (-1, -1), 0.6, colors.HexColor('#60a5fa')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
    ]
    if header:
        table_styles.extend([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0f2ff')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('FONTNAME', (0, 0), (-1, 0), PDF_FONT_BOLD),
            ('FONTNAME', (0, 1), (-1, -1), PDF_FONT_DEFAULT),
        ])
    else:
        table_styles.append(('FONTNAME', (0, 0), (-1, -1), PDF_FONT_DEFAULT))

    table.setStyle(TableStyle(table_styles))
    return table


def _render_pdf_with_reportlab(title: str, blocks: list, generated_on: str) -> io.BytesIO:
    buffer = io.BytesIO()
    _register_pdf_fonts()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
        title=title,
    )
    base_styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        'Body',
        parent=base_styles['BodyText'],
        fontName=PDF_FONT_DEFAULT,
        fontSize=11,
        leading=16,
        textColor=colors.HexColor('#1f2933'),
        spaceAfter=6,
    )
    title_style = ParagraphStyle(
        'DocumentTitle',
        parent=base_styles['Heading1'],
        fontName=PDF_FONT_BOLD,
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#005b99'),
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        'DocumentSubtitle',
        parent=base_styles['Normal'],
        fontName=PDF_FONT_DEFAULT,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#475569'),
        spaceAfter=12,
    )
    heading_styles = {
        1: ParagraphStyle('HeadingLevel1', parent=base_styles['Heading2'], fontName=PDF_FONT_BOLD, fontSize=18, leading=22, textColor=colors.HexColor('#0f172a'), spaceBefore=12, spaceAfter=6),
        2: ParagraphStyle('HeadingLevel2', parent=base_styles['Heading3'], fontName=PDF_FONT_BOLD, fontSize=15, leading=20, textColor=colors.HexColor('#0f172a'), spaceBefore=10, spaceAfter=4),
        3: ParagraphStyle('HeadingLevel3', parent=base_styles['Heading4'], fontName=PDF_FONT_BOLD, fontSize=13, leading=18, textColor=colors.HexColor('#0f172a'), spaceBefore=8, spaceAfter=3),
        4: ParagraphStyle('HeadingLevel4', parent=base_styles['Heading5'], fontName=PDF_FONT_BOLD, fontSize=12, leading=16, textColor=colors.HexColor('#0f172a'), spaceBefore=6, spaceAfter=3),
    }
    list_item_style = ParagraphStyle(
        'ListItem',
        parent=body_style,
        leftIndent=18,
        bulletIndent=18,
        spaceAfter=4,
    )
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=body_style,
        fontName=PDF_FONT_BOLD,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=2,
    )
    table_body_style = ParagraphStyle(
        'TableBody',
        parent=body_style,
        spaceAfter=0,
    )

    safe_title_text = _ensure_text(title or 'Generated Document')
    title_style_resolved = _style_for_text(title_style, safe_title_text, force_bold=True)
    subtitle_style_resolved = _style_for_text(subtitle_style, generated_on)

    story = [
        Paragraph(html.escape(safe_title_text), title_style_resolved),
        Paragraph(f'<b>Generated on:</b> {html.escape(generated_on)}', subtitle_style_resolved),
        Spacer(1, 8),
    ]

    list_buffer: list[str] = []
    list_kind: str | None = None

    def flush_list():
        nonlocal list_buffer, list_kind
        if list_buffer:
            story.extend(_build_list_flowable(list_buffer, list_kind or 'bullet', list_item_style))
            list_buffer = []
            list_kind = None

    for block in blocks:
        block_type = block.get('type')
        try:
            if block_type in ('bullet', 'numbered'):
                incoming_kind = 'bullet' if block_type == 'bullet' else 'numbered'
                if list_kind and list_kind != incoming_kind:
                    flush_list()
                list_kind = incoming_kind
                list_buffer.append(_ensure_text(block.get('text', '')))
                continue

            flush_list()

            if block_type == 'heading':
                level_value = block.get('level', 1)
                try:
                    level = int(float(level_value))
                except (TypeError, ValueError):
                    level = 1
                level = max(1, min(level, 4))
                heading_text = _ensure_text(block.get('text', ''))
                heading_style = _style_for_text(heading_styles[level], heading_text, force_bold=True)
                story.append(Paragraph(_render_inline_html(heading_text), heading_style))
            elif block_type == 'paragraph':
                paragraph_text = _ensure_text(block.get('text', ''))
                paragraph_style = _style_for_text(body_style, paragraph_text)
                story.append(Paragraph(_render_inline_html(paragraph_text), paragraph_style))
            elif block_type == 'table':
                story.append(_build_table_flowable(block, table_header_style, table_body_style))
                story.append(Spacer(1, 6))
            elif block_type == 'blank':
                story.append(Spacer(1, 8))
        except Exception as block_err:
            print(f"⚠️ Failed to render block type '{block_type}': {block_err}. Block contents: {block}")
            fallback_text = _ensure_text(block.get('text', ''))
            fallback_style = _style_for_text(body_style, fallback_text)
            story.append(Paragraph(_render_inline_html(fallback_text), fallback_style))

    flush_list()

    doc.build(story)
    buffer.seek(0)
    return buffer

def _normalize_column_key(name: str) -> str:
    """Create a comparable key for loose column matching."""
    return re.sub(r'[^a-z0-9]', '', str(name).lower())

def _find_best_column(columns, candidates):
    """Find the first column that matches any candidate alias."""
    if not columns:
        return None
    normalized_lookup = {_normalize_column_key(col): col for col in columns}
    candidate_keys = [_normalize_column_key(alias) for alias in candidates if alias]
    
    # Exact normalized match
    for candidate in candidate_keys:
        if candidate in normalized_lookup:
            return normalized_lookup[candidate]
    
    # Partial match (candidate contained within column key)
    for candidate in candidate_keys:
        for norm_key, original in normalized_lookup.items():
            if candidate and candidate in norm_key:
                return original
    
    return None

def _status_to_flag(value: str) -> str:
    """Map status strings to a normalized correctness flag."""
    text = str(value).strip().lower()
    if not text:
        return "unknown"
    correct_tokens = {"correct", "right", "true", "pass", "passed", "success", "yes", "1"}
    incorrect_tokens = {"incorrect", "wrong", "false", "fail", "failed", "no", "0"}
    if text in correct_tokens:
        return "correct"
    if text in incorrect_tokens:
        return "incorrect"
    # Heuristic: look for token presence
    if "correct" in text or "right" in text or text.startswith("y"):
        return "correct"
    if "incorrect" in text or "wrong" in text or text.startswith("n") or "fail" in text:
        return "incorrect"
    return "unknown"

print("Initializing Pinecone...")
pc = Pinecone(api_key=PINECONE_API_KEY)
pinecone_index = pc.Index(PINECONE_INDEX_NAME)
print("✅ Pinecone index connection established.")

print(f"Loading metadata from gs://{BUCKET_NAME}/{METADATA_FILE_PATH}...")
try:
    bucket = storage_client.bucket(BUCKET_NAME)
    blob = bucket.blob(METADATA_FILE_PATH)
    metadata_lookup = json.loads(blob.download_as_string())
    print("✅ Metadata loaded successfully.")
except Exception as e:
    print(f"❌ FATAL: Could not load metadata file. Error: {e}")
    metadata_lookup = {}

# --- LOCATION DETECTION ---
def detect_location_from_ip(ip_address: str) -> dict:
    """Detect location from IP address"""
    try:
        response = requests.get(f"http://ip-api.com/json/{ip_address}")
        data = response.json()
        return {
            "city": data.get("city", ""),
            "state": data.get("regionName", ""),
            "country": data.get("country", "India"),
            "detected": True
        }
    except:
        return {"city": "", "state": "", "country": "India", "detected": False}


def reverse_geocode(lat: float, lon: float) -> dict:
    """Reverse geocode latitude/longitude into a location dict."""
    try:
        headers = {"User-Agent": "TataStrive-AgentToolkit/1.0"}
        params = {
            "format": "json",
            "lat": float(lat),
            "lon": float(lon),
            "zoom": 10,
        }
        response = requests.get(
            "https://nominatim.openstreetmap.org/reverse",
            params=params,
            headers=headers,
            timeout=10,
        )
        response.raise_for_status()
        payload = response.json()
        address = payload.get("address", {})
        return {
            "city": address.get("city") or address.get("town") or address.get("village") or "",
            "state": address.get("state") or "",
            "country": address.get("country") or "India",
            "detected": True,
        }
    except Exception as exc:
        print(f"⚠️ Reverse geocode failed: {exc}")
        return {"city": "", "state": "", "country": "India", "detected": False}

def get_suggested_language(state: str) -> str:
    """Suggest language based on state"""
    language_map = {
        "West Bengal": "Bengali",
        "Maharashtra": "Marathi",
        "Gujarat": "Gujarati",
        "Tamil Nadu": "Tamil",
        "Karnataka": "Kannada",
        "Kerala": "Malayalam",
        "Andhra Pradesh": "Telugu",
        "Telangana": "Telugu",
        "Odisha": "Odia",
        "Punjab": "Punjabi",
        "Haryana": "Hindi",
        "Rajasthan": "Hindi",
        "Uttar Pradesh": "Hindi",
        "Madhya Pradesh": "Hindi",
        "Bihar": "Hindi"
    }
    return language_map.get(state, "English")

# --- DATA LOADING FUNCTIONS ---
def load_course_data(file):
    """Load course duration data from CSV"""
    global course_data
    try:
        df = pd.read_csv(file)
        if df.empty:
            raise ValueError("Course CSV is empty")

        # Normalise column names so we can support multiple schema variants
        normalized_columns = {col.lower().strip(): col for col in df.columns}
        course_name_col = normalized_columns.get("name") or normalized_columns.get("course_name")
        course_id_col = normalized_columns.get("id") or normalized_columns.get("course_id")
        duration_col = normalized_columns.get("cumulative_course_duration") or normalized_columns.get("time_in_hr")
        theory_col = normalized_columns.get("domain_theory_hours")
        eligibility_col = normalized_columns.get("eligibility_criteria")
        legend_col = normalized_columns.get("legend")

        if not course_name_col or not duration_col:
            raise ValueError(
                "Course CSV missing required columns (expected name/course_name and cumulative_course_duration/time_in_Hr)"
            )

        # Reset existing cache
        course_data.clear()

        if "cumulative_course_duration" in normalized_columns:
            # Original compact schema: one row per course with pre-computed aggregates
            for _, row in df.iterrows():
                course_name = str(row[course_name_col]).strip()
                if not course_name:
                    continue

                course_data[course_name] = {
                    "id": row[course_id_col] if course_id_col else None,
                    "duration_hours": float(row.get(duration_col, 0) or 0),
                    "theory_hours": float(row.get(theory_col, 0) or 0) if theory_col else 0.0,
                    "eligibility": row.get(eligibility_col, "Not specified") if eligibility_col else "Not specified",
                }
        else:
            # BigQuery export schema: multiple rows per course with per-session hours
            df[duration_col] = pd.to_numeric(df[duration_col], errors="coerce").fillna(0.0)

            grouped = df.groupby(course_name_col)
            for course_name, group in grouped:
                clean_name = str(course_name).strip()
                if not clean_name:
                    continue

                total_hours = float(group[duration_col].sum())
                if legend_col and legend_col in group:
                    theory_hours = float(
                        group[group[legend_col].str.contains("theory", case=False, na=False)][duration_col].sum()
                    )
                else:
                    theory_hours = 0.0

                course_id_value = group[course_id_col].iloc[0] if course_id_col and course_id_col in group else None
                eligibility_value = (
                    group[eligibility_col].dropna().iloc[0]
                    if eligibility_col and eligibility_col in group.columns and not group[eligibility_col].dropna().empty
                    else "Not specified"
                )

                course_data[clean_name] = {
                    "id": course_id_value,
                    "duration_hours": round(total_hours, 2),
                    "theory_hours": round(theory_hours, 2),
                    "eligibility": eligibility_value,
                }

        print(f"✅ Loaded {len(course_data)} courses")
        return {"success": True, "courses_loaded": len(course_data)}
    except Exception as e:
        print(f"❌ Error loading course data: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def load_holiday_data(file):
    """Load holiday data from CSV"""
    global holiday_data
    try:
        df = pd.read_csv(file)
        if df.empty:
            raise ValueError("Holiday CSV is empty")

        normalized_columns = {col.lower().strip(): col for col in df.columns}
        location_col = (
            normalized_columns.get("location")
            or normalized_columns.get("state")
            or normalized_columns.get("region")
        )
        name_col = normalized_columns.get("holidays") or normalized_columns.get("holiday")
        date_col = normalized_columns.get("holidaydate") or normalized_columns.get("date")
        day_col = normalized_columns.get("holidayday") or normalized_columns.get("day")

        if not location_col or not date_col or not name_col:
            raise ValueError("Holiday CSV missing required columns (location/state, holiday name, holiday date)")

        # Parse dates robustly (supports ISO, dd-mm-yyyy, etc.)
        parsed_dates = pd.to_datetime(df[date_col], errors="coerce", dayfirst=True)
        if parsed_dates.isna().all():
            parsed_dates = pd.to_datetime(df[date_col], errors="coerce", dayfirst=False)

        df = df.assign(_parsed_date=parsed_dates).dropna(subset=["_parsed_date"])
        holiday_data.clear()

        for state, group in df.groupby(location_col):
            state_name = str(state).strip()
            if not state_name:
                continue

            holidays_for_state = []
            for _, row in group.iterrows():
                holiday_date = row["_parsed_date"].to_pydatetime()
                holiday_day = (
                    str(row[day_col]).strip()
                    if day_col and day_col in row and pd.notna(row[day_col])
                    else holiday_date.strftime("%A")
                )
                holidays_for_state.append(
                    {
                        "name": str(row[name_col]).strip(),
                        "date": holiday_date,
                        "day": holiday_day,
                    }
                )

            if holidays_for_state:
                holiday_data[state_name] = holidays_for_state

        print(f"✅ Loaded holidays for {len(holiday_data)} states")
        return {"success": True, "states_loaded": len(holiday_data)}
    except Exception as e:
        print(f"❌ Error loading holiday data: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def load_assessment_guidelines(content):
    """Load assessment guidelines from text content"""
    global assessment_guidelines
    try:
        if isinstance(content, bytes):
            assessment_guidelines = content.decode('utf-8')
        else:
            assessment_guidelines = str(content)
        
        print(f"✅ Loaded assessment guidelines ({len(assessment_guidelines)} chars)")
        return {"success": True, "guidelines_length": len(assessment_guidelines)}
    except Exception as e:
        print(f"❌ Error loading guidelines: {e}")
        traceback.print_exc()
        return {"success": False, "error": str(e)}

def get_holidays_between_dates(state: str, start_date: datetime, end_date: datetime) -> list:
    """Get holidays for a state between two dates"""
    if state not in holiday_data:
        return []
    
    holidays = []
    for holiday in holiday_data[state]:
        if start_date <= holiday['date'] <= end_date:
            holidays.append(holiday)
    return holidays

# --- DOCUMENT DISCOVERY ---
@app.route('/get_documents', methods=['GET'])
def get_documents():
    """Fetch all unique document titles from Pinecone"""
    try:
        # Query a sample of vectors to get all unique titles
        sample_results = pinecone_index.query(
            vector=[0.0] * 768,  # Dummy vector
            top_k=1055,  # Get all records
            include_metadata=True
        )
        
        titles = set()
        for match in sample_results.get("matches", []):
            if 'metadata' in match and 'title' in match['metadata']:
                titles.add(match['metadata']['title'])
        
        sorted_titles = sorted(list(titles))
        print(f"✅ Found {len(sorted_titles)} unique documents")
        
        return jsonify({
            "documents": sorted_titles,
            "total_count": len(sorted_titles)
        }), 200
        
    except Exception as e:
        print(f"❌ Error fetching documents: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/create/content', methods=['POST'])
def create_content():
    """Create educational content based on a topic"""
    try:
        data = request.json or {}
        topic = data.get('query', '').strip()
        content_type = data.get('content_type', 'Learning Guide').strip() or "Learning Guide"
        audience = data.get('audience', 'Front Desk Associate trainees').strip() or "Front Desk Associate trainees"
        tone = data.get('tone', 'Professional').strip() or "Professional"
        length = data.get('length', 'Medium').strip() or "Medium"
        include_practice = bool(data.get('include_practice', True))
        language = data.get('language', 'English')
        format_type = data.get('output_format', 'json')
        selected_documents = data.get('selected_documents', [])
        
        if not topic:
            return jsonify({"error": "Topic is required"}), 400
        
        if not selected_documents:
            return jsonify({"error": "Please select at least one source document"}), 400
        
        context, sources = perform_rag(topic, num_neighbors=10, selected_documents=selected_documents)
        
        prompt = CONTENT_PROMPT.format(
            audience=audience,
            tone=tone,
            length_description=length,
            content_type=content_type,
            include_activity="Yes" if include_practice else "No",
            context=context
        )
        
        user_instruction = (
            f"Topic: {topic}\n"
            f"Audience profile: {audience}\n"
            f"Tone: {tone}\n"
            f"Desired content type: {content_type}\n"
            f"Preferred depth: {length}\n"
            f"Include activity or checklist: {'Yes' if include_practice else 'No'}"
        )
        
        english_content = call_openrouter(prompt, user_instruction)
        translated_content = translate_text(english_content, language)
        
        sanitized_slug = re.sub(r'[^A-Za-z0-9]+', '_', topic).strip('_') or "content"
        document_title = f"{content_type}: {topic}"
        
        if format_type == 'docx':
            doc_buffer = generate_docx(translated_content, document_title)
            return send_file(
                doc_buffer,
                as_attachment=True,
                download_name=f"content_{sanitized_slug}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        if format_type == 'pdf':
            pdf_buffer = generate_pdf(translated_content, document_title)
            return send_file(
                pdf_buffer,
                as_attachment=True,
                download_name=f"content_{sanitized_slug}.pdf",
                mimetype='application/pdf'
            )
        
        return jsonify({
            "english_answer": english_content,
            "translated_answer": translated_content,
            "language": language,
            "sources": sources,
            "metadata": {
                "content_type": content_type,
                "audience": audience,
                "tone": tone,
                "length": length,
                "include_practice": include_practice
            }
        }), 200
        
    except Exception as e:
        print(f"Error in create_content: {str(e)}")
        return jsonify({"error": str(e)}), 500

# --- CORE LOGIC ---
def perform_rag(query: str, num_neighbors: int = 5, selected_documents: list = None) -> tuple[str, list]:
    """Perform RAG search using Pinecone with optional document filtering"""
    query_embedding = embedding_model.get_embeddings([query])[0].values
    
    try:
        # Build filter if documents are specified
        query_filter = None
        if selected_documents and len(selected_documents) > 0:
            # Pinecone filter format: {"title": {"$in": ["doc1.pdf", "doc2.pdf"]}}
            query_filter = {"title": {"$in": selected_documents}}
            print(f"🔍 Filtering RAG by documents: {selected_documents}")
        
        query_results = pinecone_index.query(
            vector=query_embedding,
            top_k=num_neighbors,
            include_metadata=True,
            filter=query_filter
        )
        matches = query_results.get("matches", [])
        
        if not matches and selected_documents:
            print(f"⚠️ No matches found for selected documents. Try increasing num_neighbors or check document names.")
        
    except Exception as e:
        print(f"❌ ERROR querying Pinecone: {e}")
        traceback.print_exc()
        return "ERROR: Could not query the vector database.", []
    
    context = ""
    sources = []
    for match in matches:
        if 'metadata' in match:
            context += match['metadata'].get('text', '') + "\n---\n"
            sources.append(match['metadata'].get('title', 'Unknown'))
    
    if not context:
        context = "No relevant content found for the selected documents and query."
            
    return context, list(set(sources))

def call_openrouter(system_prompt: str, user_query: str, model: str = RAG_MODEL) -> str:
    """Call OpenRouter API for text generation with better error handling"""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://tatastrive.com",
        "X-Title": "Tata Strive RAG System"
    }
    
    data = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_query}
        ],
        "temperature": 0.3,
        "max_tokens": 3000
    }
    
    try:
        print(f"🔄 Calling OpenRouter with model: {model}")
        response = requests.post(
            f"{OPENROUTER_API_BASE}/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code != 200:
            print(f"❌ OpenRouter error: {response.status_code}")
            print(f"Response: {response.text}")
        
        response.raise_for_status()
        result = response.json()
        
        if "choices" not in result or len(result["choices"]) == 0:
            print(f"❌ No choices in response: {result}")
            return "Error: Invalid response from API"
        
        content = result["choices"][0]["message"]["content"]
        print(f"✅ OpenRouter response received ({len(content)} chars)")
        return content
        
    except requests.exceptions.Timeout:
        print(f"❌ OpenRouter timeout after 60 seconds")
        return "Error: API request timed out. Please try again."
    except requests.exceptions.HTTPError as e:
        print(f"❌ OpenRouter HTTP error: {e}")
        return f"Error: API returned error {e.response.status_code if hasattr(e, 'response') else 'unknown'}"
    except Exception as e:
        print(f"❌ OpenRouter error: {str(e)}")
        traceback.print_exc()
        return f"Error generating response: {str(e)}"

def translate_text(text: str, target_language: str) -> str:
    """Translate text to target language using OpenRouter"""
    if target_language == "English":
        return text
        
    system_prompt = f"Translate the following text to {target_language}. Maintain the original formatting and structure. Only provide the translation, no explanations."
    translated = _ensure_text(call_openrouter(system_prompt, text, TRANSLATION_MODEL)).strip()
    if not translated:
        print(f"⚠️ Translation empty for language '{target_language}'. Falling back to source content.")
        return text
    return translated

# --- PERSONALIZED LEARNING FUNCTIONS ---
def analyze_assessment_csv(df: pd.DataFrame) -> dict:
    """Analyze student assessment data"""
    try:
        if df.empty:
            raise ValueError("Assessment file is empty.")
        
        columns = list(df.columns)
        if not columns:
            raise ValueError("Assessment file has no columns.")
        
        attempt_col = _find_best_column(columns, [
            "Attempt ID", "AttemptID", "Attempt Number", "Attempt", "Submission Time", "Submitted At", "Submitted On"
        ])
        login_col = _find_best_column(columns, [
            "Login ID", "LoginID", "Email", "Student Email", "Learner Email", "User Email", "Email Address", "Username"
        ])
        question_id_col = _find_best_column(columns, [
            "Question ID", "QuestionID", "Question Number", "Question No", "Question Code"
        ])
        question_text_col = _find_best_column(columns, [
            "Question Text", "Question", "Question Statement", "Question Description", "Question Title"
        ])
        status_col = _find_best_column(columns, [
            "Answer Status", "Status", "Result", "Answer Result", "Is Correct", "Outcome", "Response Status"
        ])
        marks_col = _find_best_column(columns, [
            "Obtained Marks", "Score", "Marks Obtained", "Marks", "Points", "Earned Points", "Awarded Score"
        ])
        
        if not login_col:
            raise ValueError("Could not find a student email/login column in the assessment file.")
        
        if not question_id_col and not question_text_col:
            raise ValueError("Could not find a question identifier or question text column in the assessment file.")
        
        if not status_col and not marks_col:
            raise ValueError("Assessment file missing answer status or score information.")
        
        working_df = df.copy()
        
        # Sort by attempt or submission time to get the latest attempts
        if attempt_col and attempt_col in working_df:
            try:
                attempt_values = pd.to_datetime(working_df[attempt_col], errors='coerce')
                if attempt_values.notna().any():
                    working_df = working_df.assign(__attempt_sort=attempt_values.fillna(pd.Timestamp.min))
                    working_df = working_df.sort_values('__attempt_sort')
                else:
                    working_df = working_df.sort_values(attempt_col, na_position='last')
            except Exception:
                working_df = working_df.sort_values(attempt_col, na_position='last')
        else:
            working_df = working_df.reset_index(drop=True)
        
        group_keys = [login_col]
        question_key = question_id_col or question_text_col
        if question_key:
            group_keys.append(question_key)
        
        if group_keys:
            working_df = working_df.groupby(group_keys, as_index=False).tail(1)
        
        # Prepare marks and status flags
        if marks_col and marks_col in working_df:
            working_df['__marks'] = pd.to_numeric(working_df[marks_col], errors='coerce').fillna(0.0)
        else:
            if not status_col:
                raise ValueError("Unable to determine marks because status column is missing.")
            working_df['__marks'] = working_df[status_col].apply(
                lambda v: 1.0 if _status_to_flag(v) == "correct" else 0.0
            )
        
        if status_col and status_col in working_df:
            working_df['__status_norm'] = working_df[status_col].apply(_status_to_flag)
        else:
            working_df['__status_norm'] = working_df['__marks'].apply(lambda m: "correct" if m > 0 else "incorrect")
        
        # Aggregate student performance
        student_performance = working_df.groupby(login_col).agg(
            total_marks=pd.NamedAgg(column='__marks', aggfunc='sum'),
            total_questions=pd.NamedAgg(column='__marks', aggfunc='count')
        ).reset_index()
        student_performance = student_performance.rename(columns={login_col: 'student_id'})
        student_performance['percentage'] = (
            (student_performance['total_marks'] / student_performance['total_questions']).replace([float('inf'), -float('inf')], 0) * 100
        )
        student_performance['percentage'] = student_performance['percentage'].fillna(0.0)
        
        weak_students_df = student_performance[student_performance['percentage'] < 70]
        weak_students = weak_students_df.to_dict('records')
        
        question_display_col = question_text_col or question_id_col
        
        student_details = []
        for student in weak_students:
            student_email = student['student_id']
            failed_rows = working_df[
                (working_df[login_col] == student_email) &
                (working_df['__status_norm'] != "correct")
            ]
            
            if failed_rows.empty:
                continue
            
            failed_questions = failed_rows[question_display_col].astype(str).tolist()
            student_details.append({
                "email": student_email,
                "score": f"{int(student['total_marks'])}/{int(student['total_questions'])}",
                "percentage": round(student['percentage'], 1),
                "failed_questions": failed_questions
            })
        
        question_performance = working_df.groupby(question_display_col).agg(
            success_rate=pd.NamedAgg(
                column='__status_norm',
                aggfunc=lambda x: (x.eq("correct").sum() / len(x)) * 100 if len(x) else 0
            )
        ).reset_index()
        question_performance = question_performance.rename(columns={question_display_col: 'question'})
        weak_questions = question_performance[question_performance['success_rate'] < 60].to_dict('records')
        
        return {
            "total_students": int(student_performance['student_id'].nunique()),
            "weak_students": weak_students,
            "weak_questions": weak_questions,
            "average_score": float(student_performance['percentage'].mean() if not student_performance.empty else 0.0),
            "student_details": student_details
        }
    except Exception as e:
        print(f"Error analyzing assessment: {e}")
        traceback.print_exc()
        return {"error": str(e)}

def extract_topic_from_question(question_text: str) -> str:
    """Extract key topic from question text for RAG search"""
    clean_text = re.sub(r'<[^>]+>', '', question_text)
    clean_text = clean_text.replace('<DOUBLE_QUOTES>', '"')
    clean_text = clean_text.replace('<COMMA>', ',')
    clean_text = clean_text.replace('<br>', ' ')
    clean_text = re.sub(r'Select the correct option\..*$', '', clean_text)
    clean_text = re.sub(r'\?.*$', '', clean_text)
    
    keywords_map = {
        'hotel': 'hotel definition and types',
        'restaurant': 'restaurant and food service',
        'hospitality': 'hospitality industry basics',
        'sarai': 'traditional accommodation types in India',
        'dharamshala': 'traditional accommodation types in India',
        'front office': 'front office department and management',
        'city hotels': 'hotel classifications and types',
        'thinnai': 'traditional Indian hospitality culture',
        'independence': 'history of hotel industry in India',
        'boat houses': 'resort hotels and specialized accommodations',
        'dal lake': 'resort hotels and specialized accommodations'
    }
    
    clean_lower = clean_text.lower()
    for keyword, topic in keywords_map.items():
        if keyword in clean_lower:
            return f"Front Desk Associate {topic}"
    
    return "Front Desk Associate hospitality basics"

def clean_question_text(question_text: str) -> str:
    """Clean question text for display"""
    clean = re.sub(r'<[^>]+>', '', question_text)
    clean = clean.replace('<DOUBLE_QUOTES>', '"')
    clean = clean.replace('<COMMA>', ',')
    clean = clean.replace('<br>', ' ')
    if len(clean) > 100:
        clean = clean[:100] + "..."
    return clean

def generate_personalized_content_for_student(student_email: str, failed_questions: list) -> str:
    """Generate personalized study content for a student"""
    topics = [extract_topic_from_question(q) for q in failed_questions]
    weak_topics = ", ".join(set(topics))
    
    all_context = ""
    for topic in set(topics):
        context, _ = perform_rag(topic, num_neighbors=3)
        all_context += context + "\n\n"
    
    prompt = f"""You are an expert educator creating personalized remedial content for a Tata Strive Front Desk Associate student.

Student: {student_email}
Topics they struggled with: {weak_topics}

Questions they got wrong:
{chr(10).join([f"- {q}" for q in failed_questions])}

Context from knowledge base:
{all_context}

Create a comprehensive study guide that:
1. Explains each concept in simple, clear language
2. Addresses common misconceptions
3. Provides practical examples relevant to Front Desk Associate work
4. Includes memory aids and tips
5. Adds 5 practice questions with detailed explanations
6. Uses encouraging, supportive tone

Focus on helping them understand these specific topics better."""
    
    content = call_openrouter(prompt, "Generate personalized study guide")
    return content

def send_email_with_pdf(to_email: str, subject: str, body_plain: str, body_html: str, pdf_content: bytes, pdf_name: str) -> bool:
    """Send multipart email with PDF attachment using Gmail SMTP"""
    try:
        msg = MIMEMultipart('alternative')
        msg['From'] = f"{EMAIL_SENDER_NAME} <{EMAIL_SENDER}>"
        msg['To'] = to_email
        msg['Subject'] = subject
        
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #2c3e50;">Hi Student,</h2>
                
                <p>You recently completed the <strong>Front Desk Associate</strong> assessment.</p>
                
                {body_html}
                
                <div style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 12px;">
                    <p>This is an automated message from Tata Strive Learning Platform.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(body_plain, 'plain'))
        msg.attach(MIMEText(html_body, 'html'))
        
        pdf_attachment = MIMEApplication(pdf_content, _subtype="pdf")
        pdf_attachment.add_header('Content-Disposition', 'attachment', filename=pdf_name)
        msg.attach(pdf_attachment)
        
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.send_message(msg)
        
        print(f"✅ Email sent successfully to {to_email}")
        return True
        
    except Exception as e:
        print(f"❌ Failed to send email to {to_email}: {str(e)}")
        traceback.print_exc()
        return False

# --- PROMPTS ---
ASSESSMENT_PROMPT = """You are an expert educational assessment creator for Tata Strive programs.

GUIDELINES TO FOLLOW:
{guidelines}

Context from knowledge base:
{context}

Create an assessment that:
1. Follows ALL guidelines strictly (MCSS/MCMS format, 4 options, no "All of the above")
2. Uses Bloom's taxonomy levels 1-2 (Knowledge and Comprehension)
3. Tests practical understanding
4. Provides clear marking criteria
5. Is culturally sensitive and uses inclusive language

Format the output clearly with numbered questions and marking schemes."""

LESSON_PLAN_PROMPT = """You are an expert curriculum designer for Tata Strive programs.

Course Information:
{course_info}

Holidays to Consider:
{holidays}

Context from knowledge base:
{context}

Create a detailed lesson plan that:
1. Spans the exact course duration ({duration} hours)
2. EXCLUDES all holidays listed above
3. EXCLUDES weekends (Saturday/Sunday)
4. Includes:
   - Learning objectives
   - Daily schedule with time allocation
   - Module-wise breakdown
   - Assessment strategies
   - Materials needed
   - Key takeaways

Make it practical, realistic, and aligned with Tata Strive's teaching methodology.
Ensure the timeline accounts for holidays and weekends."""

CONTENT_PROMPT = """You are an expert educational content creator supporting Tata Strive programs.

Audience: {audience}
Preferred tone: {tone}
Desired depth: {length_description}
Requested format: {content_type}

Context from knowledge base:
{context}

Create content that:
1. Opens with a short summary and clear learning objectives.
2. Uses concise sections with helpful subheadings.
3. Includes practical, India-relevant examples and scenarios.
4. Provides at least one actionable activity or checklist if {include_activity}.
5. Ends with key takeaways and next steps for facilitators.

Keep the language encouraging, inclusive, and aligned with Tata Strive's methodology."""

# --- DOCUMENT GENERATION ---
def _parse_structured_content(content: str) -> list:
    """Convert Markdown-like text into structured blocks for export."""
    text_content = _ensure_text(content)
    if not text_content.strip():
        return [{"type": "paragraph", "text": "Content could not be generated. Please try again."}]
    
    lines = text_content.replace('\r\n', '\n').replace('\r', '\n').split('\n')
    blocks = []
    table_buffer = []
    
    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            blocks.append(_build_table_block(table_buffer))
            table_buffer = []
    
    heading_pattern = re.compile(r'^(#{1,6})\s+(.*)$')
    numbered_pattern = re.compile(r'^(\d+)[\.\)]\s+(.*)$')
    
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        
        if table_buffer and not (stripped.startswith("|") and '|' in stripped[1:]):
            flush_table()
        
        if not stripped:
            if blocks and blocks[-1].get("type") == "blank":
                continue
            blocks.append({"type": "blank"})
            continue
        
        if stripped.startswith("|") and '|' in stripped[1:]:
            table_buffer.append(stripped)
            continue
        
        heading_match = heading_pattern.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            continue
        
        if stripped.startswith(("- ", "* ")):
            text = stripped[2:].strip()
            blocks.append({"type": "bullet", "text": text})
            continue
        
        numbered_match = numbered_pattern.match(stripped)
        if numbered_match:
            text = numbered_match.group(2).strip()
            blocks.append({"type": "numbered", "text": text})
            continue
        
        blocks.append({"type": "paragraph", "text": stripped})
    
    flush_table()
    return blocks


def _build_table_block(lines: list) -> dict:
    """Interpret a markdown style table block."""
    rows = []
    for line in lines:
        clean = line.strip().strip('|')
        cells = [cell.strip() for cell in clean.split('|')]
        rows.append(cells)
    
    header = []
    body = rows
    if len(rows) >= 2:
        divider_candidates = rows[1]
        if all(re.fullmatch(r':?-{3,}:?', cell.replace(' ', '')) for cell in divider_candidates):
            header = rows[0]
            body = rows[2:]

    col_count = len(header) if header else (len(body[0]) if body else 0)
    return {"type": "table", "header": header, "rows": body, "cols": col_count}


def _clear_paragraph(paragraph) -> None:
    """Remove all existing runs from a python-docx paragraph."""
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _add_markdown_runs(paragraph, text: str, *, default_bold: bool = False) -> None:
    """Render lightweight markdown emphasis into docx runs without leaving raw markers."""
    content = html.unescape(_ensure_text(text))
    if not content:
        return

    bold_state = default_bold
    bold_toggle = False
    italic_state = False
    code_state = False
    buffer: list[str] = []
    text_length = len(content)
    i = 0

    def flush_buffer():
        nonlocal buffer
        if not buffer:
            return
        run = paragraph.add_run(''.join(buffer))
        run.bold = bold_state
        run.italic = italic_state
        if code_state:
            run.font.name = "Courier New"
        buffer = []

    while i < text_length:
        segment = content[i:]
        if segment.startswith("**") or segment.startswith("__"):
            flush_buffer()
            bold_toggle = not bold_toggle
            bold_state = default_bold ^ bold_toggle
            i += 2
            continue
        if segment.startswith("*") or segment.startswith("_"):
            flush_buffer()
            italic_state = not italic_state
            i += 1
            continue
        if segment.startswith("`"):
            flush_buffer()
            code_state = not code_state
            i += 1
            continue
        if segment.startswith("\\") and len(segment) > 1:
            buffer.append(segment[1])
            i += 2
            continue
        if segment.startswith("\n"):
            flush_buffer()
            paragraph.add_run().add_break()
            i += 1
            continue
        buffer.append(segment[0])
        i += 1

    flush_buffer()

    if not paragraph.runs:
        run = paragraph.add_run("")
        run.bold = default_bold


def generate_docx(content: str, title: str) -> io.BytesIO:
    """Generate a styled DOCX document from structured content."""
    doc = Document()
    cover_heading = doc.add_heading(level=0)
    _clear_paragraph(cover_heading)
    _add_markdown_runs(cover_heading, title or "Generated Document", default_bold=True)
    
    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
    except Exception:
        pass
    
    blocks = _parse_structured_content(str(content))
    
    for block in blocks:
        block_type = block.get("type")
        if block_type == "heading":
            level_value = block.get("level", 1)
            try:
                level = int(float(level_value))
            except (TypeError, ValueError):
                level = 1
            level = max(1, min(level, 4))
            heading_para = doc.add_heading(level=level)
            _clear_paragraph(heading_para)
            _add_markdown_runs(heading_para, block.get("text", ""), default_bold=True)
        elif block_type == "paragraph":
            para = doc.add_paragraph()
            _clear_paragraph(para)
            _add_markdown_runs(para, block.get("text", ""))
        elif block_type == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            _clear_paragraph(para)
            _add_markdown_runs(para, block.get("text", ""))
        elif block_type == "numbered":
            para = doc.add_paragraph(style="List Number")
            _clear_paragraph(para)
            _add_markdown_runs(para, block.get("text", ""))
        elif block_type == "table":
            header = [_ensure_text(cell) for cell in block.get("header", [])]
            rows = [
                [_ensure_text(cell) for cell in (_row or [])]
                for _row in block.get("rows", [])
            ]
            cols_value = block.get("cols")
            try:
                cols = int(cols_value)
            except (TypeError, ValueError):
                cols = 0
            if cols <= 0:
                cols = len(header) if header else len(rows[0]) if rows else 0
            if cols <= 0:
                continue
            total_rows = (1 if header else 0) + len(rows)
            table = doc.add_table(rows=total_rows, cols=cols)
            try:
                table.style = "Light Shading Accent 1"
            except Exception:
                pass
            row_idx = 0
            if header:
                hdr_cells = table.rows[0].cells
                for idx in range(min(cols, len(header))):
                    paragraph = hdr_cells[idx].paragraphs[0]
                    _clear_paragraph(paragraph)
                    _add_markdown_runs(paragraph, header[idx], default_bold=True)
                row_idx = 1
            for data_row in rows:
                cells = table.rows[row_idx].cells
                for idx in range(cols):
                    text = data_row[idx] if idx < len(data_row) else ""
                    paragraph = cells[idx].paragraphs[0]
                    _clear_paragraph(paragraph)
                    _add_markdown_runs(paragraph, text)
                row_idx += 1
            doc.add_paragraph("")
        elif block_type == "blank":
            doc.add_paragraph("")
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


def generate_pdf(content: str, title: str) -> io.BytesIO:
    """Generate a polished PDF; fallback to ReportLab when WeasyPrint is unavailable."""
    fallback_message = "Content could not be generated. Please try again."
    raw_text = _ensure_text(content)
    blocks = _parse_structured_content(raw_text)

    has_substantive_content = any(
        block.get("type") not in {"blank"} and bool(_ensure_text(block.get("text", "")).strip())
        for block in blocks
    )

    if not has_substantive_content:
        markdown_text = fallback_message
        blocks = _parse_structured_content(fallback_message)
    else:
        markdown_text = raw_text

    generated_on = datetime.now().strftime("%d %B %Y")
    chosen_title = title or "Generated Document"

    if WEASYPRINT_AVAILABLE and HTML and CSS:
        body_html = markdown_renderer.render(markdown_text)
        safe_title = html.escape(chosen_title)
        document_html = f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <title>{safe_title}</title>
  </head>
  <body>
    <section class="title-header">
      <h1>{safe_title}</h1>
      <p class="subtitle">Generated on {generated_on}</p>
    </section>
    <article class="content">
      {body_html}
    </article>
  </body>
</html>
"""
        pdf_buffer = io.BytesIO()
        HTML(string=document_html, base_url=os.getcwd()).write_pdf(
            pdf_buffer,
            stylesheets=[CSS(string=PDF_BASE_STYLE_CSS)]
        )
        pdf_buffer.seek(0)
        return pdf_buffer

    global WEASYPRINT_WARNING_EMITTED
    if not WEASYPRINT_WARNING_EMITTED:
        print("WeasyPrint not available; using ReportLab renderer for PDFs.")
        if WEASYPRINT_IMPORT_ERROR:
            print(f"WeasyPrint import error: {WEASYPRINT_IMPORT_ERROR}")
        WEASYPRINT_WARNING_EMITTED = True

    try:
        return _render_pdf_with_reportlab(chosen_title, blocks, generated_on)
    except Exception as err:
        print(f"⚠️ ReportLab styled PDF renderer failed: {err}")
        traceback.print_exc()
        return _render_plain_pdf_fallback(chosen_title, generated_on, raw_text)


def _render_plain_pdf_fallback(title: str, generated_on: str, body_text: str) -> io.BytesIO:
    """Last-resort PDF fallback using a simple canvas to avoid hard failures."""
    buffer = io.BytesIO()
    _register_pdf_fonts()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin_x = 0.85 * inch
    margin_y = 1.0 * inch
    y = height - margin_y

    safe_title = _ensure_text(title or "Generated Document")
    safe_generated = _ensure_text(generated_on)
    safe_body = _ensure_text(body_text)

    rendered_html = markdown_renderer.render(safe_body)
    plain_body = re.sub(r"<[^>]+>", "\n", rendered_html)
    plain_body = html.unescape(plain_body)

    pdf.setFillColor(colors.HexColor('#005b99'))
    pdf.setFont(PDF_FONT_BOLD if _PDF_FONTS_REGISTERED else "Helvetica-Bold", 18)
    pdf.drawString(margin_x, y, safe_title)
    y -= 22

    pdf.setFillColor(colors.HexColor('#475569'))
    pdf.setFont(PDF_FONT_DEFAULT if _PDF_FONTS_REGISTERED else "Helvetica", 10)
    pdf.drawString(margin_x, y, f"Generated on: {safe_generated}")
    y -= 18

    pdf.setFillColor(colors.HexColor('#1f2933'))
    pdf.setFont(PDF_FONT_DEFAULT if _PDF_FONTS_REGISTERED else "Helvetica", 11)
    line_height = 14
    min_y = margin_y

    for line in plain_body.splitlines():
        if y <= min_y:
            pdf.showPage()
            y = height - margin_y
            pdf.setFillColor(colors.HexColor('#1f2933'))
            pdf.setFont(PDF_FONT_DEFAULT if _PDF_FONTS_REGISTERED else "Helvetica", 11)
        stripped = line.strip()
        if not stripped:
            y -= line_height
            continue
        pdf.drawString(margin_x, y, stripped)
        y -= line_height

    pdf.save()
    buffer.seek(0)
    return buffer




# --- ENDPOINTS ---
@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy", 
        "service": "Tata Strive RAG API",
        "courses_loaded": len(course_data),
        "states_with_holidays": len(holiday_data),
        "guidelines_loaded": len(assessment_guidelines) > 0
    }), 200

@app.route('/detect_location', methods=['POST'])
def detect_location():
    """Detect location from IP and suggest language"""
    try:
        data = request.json or {}
        lat = data.get('lat')
        lon = data.get('lon')
        location = None

        if lat is not None and lon is not None:
            try:
                location = reverse_geocode(float(lat), float(lon))
            except (TypeError, ValueError):
                location = None

        if not location or not location.get("detected"):
            ip_address = data.get('ip', request.remote_addr)
            location = detect_location_from_ip(ip_address)
        
        suggested_lang = get_suggested_language(location['state'])
        
        return jsonify({
            "location": location,
            "suggested_language": suggested_lang
        }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload/course_data', methods=['POST'])
def upload_course_data():
    """Upload course duration data"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        result = load_course_data(file)
        return jsonify(result), 200 if result['success'] else 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload/holidays', methods=['POST'])
def upload_holidays():
    """Upload holiday data"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        result = load_holiday_data(file)
        return jsonify(result), 200 if result['success'] else 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload/guidelines', methods=['POST'])
def upload_guidelines():
    """Upload assessment guidelines"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        result = load_assessment_guidelines(file.read())
        return jsonify(result), 200 if result['success'] else 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/create/assessment', methods=['POST'])
def create_assessment():
    """Create an assessment based on a topic"""
    try:
        data = request.json
        topic = data.get('query', '')
        requirements = data.get('requirements', '')
        language = data.get('language', 'English')
        format_type = data.get('output_format', 'json')
        selected_documents = data.get('selected_documents', [])
        
        if not topic:
            return jsonify({"error": "Query/topic is required"}), 400
        
        if not selected_documents or len(selected_documents) == 0:
            return jsonify({"error": "Please select at least one source document"}), 400
        
        context, sources = perform_rag(topic, num_neighbors=10, selected_documents=selected_documents)
        full_query = f"Topic: {topic}\nRequirements: {requirements}"
        
        prompt = ASSESSMENT_PROMPT.format(
            guidelines=assessment_guidelines if assessment_guidelines else "Standard MCQ format with 4 options.",
            context=context
        )
        
        assessment = call_openrouter(prompt, full_query)
        translated_assessment = translate_text(assessment, language)
        
        if format_type == 'docx':
            doc_buffer = generate_docx(translated_assessment, f"Assessment: {topic}")
            return send_file(doc_buffer, as_attachment=True, 
                           download_name=f"assessment_{topic.replace(' ', '_')}.docx",
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif format_type == 'pdf':
            pdf_buffer = generate_pdf(translated_assessment, f"Assessment: {topic}")
            return send_file(pdf_buffer, as_attachment=True,
                           download_name=f"assessment_{topic.replace(' ', '_')}.pdf",
                           mimetype='application/pdf')
        
        return jsonify({
            "english_answer": assessment,
            "translated_answer": translated_assessment,
            "language": language,
            "sources": sources
        }), 200
        
    except Exception as e:
        print(f"Error in create_assessment: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/create/lesson_plan', methods=['POST'])
def create_lesson_plan():
    """Create a lesson plan based on a topic"""
    try:
        data = request.json
        topic = data.get('query', '')
        course_name = data.get('course_name', '')
        state = data.get('state', 'Corporate')
        start_date_str = data.get('start_date', '')
        language = data.get('language', 'English')
        format_type = data.get('output_format', 'json')
        selected_documents = data.get('selected_documents', [])
        
        if not topic:
            return jsonify({"error": "Topic is required"}), 400
        
        if not selected_documents or len(selected_documents) == 0:
            return jsonify({"error": "Please select at least one source document"}), 400
        
        course_info_str = "Duration not specified"
        duration_hours = 0
        if course_name and course_name in course_data:
            course_info = course_data[course_name]
            duration_hours = course_info['duration_hours']
            course_info_str = f"Course: {course_name}\nDuration: {duration_hours} hours\nTheory Hours: {course_info['theory_hours']}\nEligibility: {course_info['eligibility']}"
        
        holidays_str = "No holidays data available"
        if state and start_date_str and state in holiday_data:
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
                days_needed = int(duration_hours / 7) if duration_hours > 0 else 30
                end_date = start_date + timedelta(days=days_needed * 2)
                
                holidays = get_holidays_between_dates(state, start_date, end_date)
                if holidays:
                    holidays_str = "Holidays to EXCLUDE:\n" + "\n".join([
                        f"- {h['name']}: {h['date'].strftime('%d %B %Y')} ({h['day']})"
                        for h in holidays
                    ])
            except:
                pass
        
        context, sources = perform_rag(topic, num_neighbors=10, selected_documents=selected_documents)
        
        full_query = f"Topic: {topic}\nCreate lesson plan considering all holidays and weekends."
        prompt = LESSON_PLAN_PROMPT.format(
            course_info=course_info_str,
            duration=f"{duration_hours}" if duration_hours > 0 else "unspecified",
            holidays=holidays_str,
            context=context
        )
        
        lesson_plan = call_openrouter(prompt, full_query)
        translated_lesson_plan = translate_text(lesson_plan, language)
        
        if format_type == 'docx':
            doc_buffer = generate_docx(translated_lesson_plan, f"Lesson Plan: {topic}")
            return send_file(doc_buffer, as_attachment=True,
                           download_name=f"lesson_plan_{topic.replace(' ', '_')}.docx",
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif format_type == 'pdf':
            pdf_buffer = generate_pdf(translated_lesson_plan, f"Lesson Plan: {topic}")
            return send_file(pdf_buffer, as_attachment=True,
                           download_name=f"lesson_plan_{topic.replace(' ', '_')}.pdf",
                           mimetype='application/pdf')
        
        return jsonify({
            "english_answer": lesson_plan,
            "translated_answer": translated_lesson_plan,
            "language": language,
            "sources": sources,
            "holidays_considered": holidays_str
        }), 200
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- PERSONALIZED LEARNING ENDPOINT ---
@app.route('/process/assessment_and_email', methods=['POST'])
def process_assessment_and_email():
    """Process assessment CSV and automatically email personalized content to weak students"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        elif file.filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file)
        else:
            return jsonify({"error": "Invalid file format. Upload CSV or Excel"}), 400
        
        print("📊 Analyzing assessment data...")
        analysis = analyze_assessment_csv(df)
        
        if 'error' in analysis:
            return jsonify({"error": analysis['error']}), 500
        
        email_results = []
        student_details = analysis.get('student_details', [])
        
        print(f"📧 Processing {len(student_details)} students who need support...")
        
        for student in student_details:
            try:
                student_email = student['email']
                failed_questions = student['failed_questions']
                score = student['score']
                percentage = student['percentage']
                
                print(f"Generating content for {student_email}...")
                
                content = generate_personalized_content_for_student(student_email, failed_questions)
                
                pdf_buffer = generate_pdf(content, f"Personalized Study Guide - {student_email}")
                pdf_bytes = pdf_buffer.getvalue()
                
                cleaned_questions = [clean_question_text(q) for q in failed_questions[:5]]
                topics_plain = "\n".join([f"- {q}" for q in cleaned_questions]) if cleaned_questions else ""
                
                body_plain = (
                    "Hi Student,\n\n"
                    "You recently completed the Front Desk Associate assessment.\n\n"
                    "YOUR RESULTS:\n"
                    f"Score: {score} ({percentage}%)\n"
                    "Status: Needs Improvement\n\n"
                    "We've analyzed your performance and attached a personalized study guide to help you revisit the areas that need attention.\n"
                    "Topics highlighted for revision:\n"
                    f"{topics_plain if topics_plain else '- Personalized to your recent attempt'}\n\n"
                    "Review the PDF before your next attempt to make the best progress.\n\n"
                    "Best regards,\n"
                    f"{EMAIL_SENDER_NAME}"
                )

                body_html = f"""
                <div style="background-color: #f8f9fa; padding: 20px; border-radius: 5px; margin: 20px 0;">
                    <h3 style="color: #2c3e50; margin-top: 0;">📊 YOUR RESULTS:</h3>
                    <p style="font-size: 18px;"><strong>Score:</strong> {score} ({percentage}%)</p>
                    <p style="color: #e74c3c;"><strong>Status:</strong> Needs Improvement</p>
                </div>
                <p>We've analyzed your performance and created a personalized study guide to help you master the concepts you found challenging.</p>
                <div style="background-color: #fff3cd; padding: 15px; border-left: 4px solid #ffc107; margin: 20px 0;">
                    <h4 style="color: #856404; margin-top: 0;">⚠️ TOPICS YOU STRUGGLED WITH:</h4>
                    <ul style="color: #856404;">
                        {chr(10).join([f"<li>{q}</li>" for q in cleaned_questions])}
                    </ul>
                </div>
                <div style="background-color: #d4edda; padding: 15px; border-radius: 5px; margin: 20px 0;">
                    <h4 style="color: #155724; margin-top: 0;">📎 ATTACHED: Your_Personalized_Study_Guide.pdf</h4>
                    <p style="color: #155724; margin-bottom: 10px;"><strong>This guide includes:</strong></p>
                    <ul style="color: #155724;">
                        <li>Clear explanations of each topic</li>
                        <li>Practical examples for Front Desk work</li>
                        <li>Memory tips and tricks</li>
                        <li>Practice questions with answers</li>
                    </ul>
                </div>
                <p style="background-color: #e3f2fd; padding: 10px; border-left: 4px solid #2196f3; margin: 20px 0;">
                    <strong>💡 TIP:</strong> Review this guide before your next attempt!
                </p>
                <p style="margin-top: 30px;">Best regards,<br><strong>{EMAIL_SENDER_NAME}</strong></p>
                """
                
                subject = "📚 Your Personalized Study Guide - Front Desk Associate"
                
                success = send_email_with_pdf(
                    to_email=student_email,
                    subject=subject,
                    body_plain=body_plain,
                    body_html=body_html,
                    pdf_content=pdf_bytes,
                    pdf_name=f"Study_Guide_{student_email.split('@')[0]}.pdf"
                )
                
                email_results.append({
                    "email": student_email,
                    "status": "✅ Sent" if success else "❌ Failed",
                    "score": score,
                    "percentage": percentage
                })
                
            except Exception as e:
                print(f"Error processing student {student['email']}: {str(e)}")
                email_results.append({
                    "email": student['email'],
                    "status": f"❌ Error: {str(e)}",
                    "score": student.get('score', 'N/A'),
                    "percentage": student.get('percentage', 'N/A')
                })
        
        return jsonify({
            "total_students": analysis['total_students'],
            "average_score": round(analysis['average_score'], 1),
            "emails_sent": len([r for r in email_results if "✅" in r['status']]),
            "email_results": email_results,
            "weak_questions": analysis['weak_questions']
        }), 200
        
    except Exception as e:
        print(f"Error in process_assessment_and_email: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/search', methods=['POST'])
def search():
    """Direct search endpoint for testing"""
    try:
        data = request.json
        query = data.get('query', '')
        selected_documents = data.get('selected_documents', None)
        
        if not query:
            return jsonify({"error": "Query is required"}), 400
        
        context, sources = perform_rag(query, num_neighbors=5, selected_documents=selected_documents)
        
        return jsonify({
            "context": context,
            "sources": sources
        }), 200
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    print("Starting Tata Strive RAG API Server...")
    print("=" * 50)
    print("🚀 Tata Strive RAG API Server")
    print("=" * 50)
    print(f"📊 Pinecone Index: {PINECONE_INDEX_NAME}")
    print(f"🤖 LLM Model: {RAG_MODEL}")
    print(f"🌐 Translation Model: {TRANSLATION_MODEL}")
    print(f"📧 Email Sender: {EMAIL_SENDER}")
    print(f"🔧 Server running on: http://0.0.0.0:8081")
    print("=" * 50)
    print("\nAvailable Endpoints:")
    print("  GET  /health")
    print("  GET  /get_documents")
    print("  POST /detect_location")
    print("  POST /upload/course_data")
    print("  POST /upload/holidays")
    print("  POST /upload/guidelines")
    print("  POST /create/assessment")
    print("  POST /create/lesson_plan")
    print("  POST /create/content")
    print("  POST /process/assessment_and_email")
    print("  POST /search")
    print("=" * 50)
    print("\n✅ Server is ready to accept requests!\n")
    
    app.run(host='0.0.0.0', port=8081, debug=True)
